from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocumentWriter:
    """Convert a markdown string into a python-docx Document and save it."""

    _H1_RE = re.compile(r"^# (.+)$")
    _H2_RE = re.compile(r"^## (.+)$")
    _H3_RE = re.compile(r"^### (.+)$")
    _BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
    _BULLET_RE = re.compile(r"^[-*] (.+)$")
    _NUMBERED_RE = re.compile(r"^\d+\. (.+)$")
    _HR_RE = re.compile(r"^---+$")

    def create(self, markdown_content: str) -> Document:
        doc = Document()
        self._set_default_style(doc)
        self._parse_and_fill(doc, markdown_content)
        return doc

    def save(self, doc: Document, path: str) -> str:
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        return path

    def open_file(self, path: str) -> None:
        os.startfile(path)

    # ------------------------------------------------------------------
    def _set_default_style(self, doc: Document) -> None:
        style = doc.styles["Normal"]
        style.font.name = "맑은 고딕"
        style.font.size = Pt(10)

    def _parse_and_fill(self, doc: Document, content: str) -> None:
        lines = content.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]

            if m := self._H1_RE.match(line):
                p = doc.add_heading(m.group(1), level=1)
                p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            elif m := self._H2_RE.match(line):
                doc.add_heading(m.group(1), level=2)
            elif m := self._H3_RE.match(line):
                doc.add_heading(m.group(1), level=3)
            elif self._HR_RE.match(line):
                doc.add_paragraph("─" * 50)
            elif m := self._BULLET_RE.match(line):
                p = doc.add_paragraph(style="List Bullet")
                self._add_run_with_bold(p, m.group(1))
            elif m := self._NUMBERED_RE.match(line):
                p = doc.add_paragraph(style="List Number")
                self._add_run_with_bold(p, m.group(1))
            elif line.startswith("|") and "|" in line[1:]:
                table_lines = [line]
                j = i + 1
                while j < len(lines) and lines[j].startswith("|"):
                    table_lines.append(lines[j])
                    j += 1
                self._add_table(doc, table_lines)
                i = j
                continue
            elif line.strip():
                p = doc.add_paragraph()
                self._add_run_with_bold(p, line)
            else:
                doc.add_paragraph()

            i += 1

    def _add_run_with_bold(self, paragraph, text: str) -> None:
        parts = self._BOLD_RE.split(text)
        for idx, part in enumerate(parts):
            run = paragraph.add_run(part)
            run.bold = (idx % 2 == 1)

    def _add_table(self, doc: Document, table_lines: list[str]) -> None:
        rows = [
            [cell.strip() for cell in line.strip("|").split("|")]
            for line in table_lines
            if not re.match(r"^\|[-| :]+\|$", line.strip())
        ]
        if not rows:
            return
        col_count = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
