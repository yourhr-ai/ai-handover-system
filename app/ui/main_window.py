import hashlib
import html
import json
import math
import re
import time
from collections import Counter
from pathlib import Path
from datetime import datetime
from dataclasses import replace

from PySide6.QtCore import QEvent, QPointF, QRectF, QThread, QTimer, Signal, QSize, Slot, Qt, QUrl
from PySide6.QtGui import QColor, QDesktopServices, QFont, QPainter, QPen
from PySide6.QtWidgets import (
    QApplication,
    QComboBox,
    QDialog,
    QDialogButtonBox,
    QFileDialog,
    QFrame,
    QGraphicsDropShadowEffect,
    QHBoxLayout,
    QInputDialog,
    QLabel,
    QLineEdit,
    QListWidget,
    QListWidgetItem,
    QMainWindow,
    QMessageBox,
    QPushButton,
    QStackedLayout,
    QTabBar,
    QTabWidget,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from app.api_config import load_api_key, save_api_key
from app.size_units import bytes_to_gb, format_gb
from app.license import (
    check_server_reachable,
    get_device_id,
    is_license_active,
    load_saved_license_code,
    save_license,
    validate_license,
    verify_license_with_server,
)
from app.license_credits import (
    check_balance,
    consume_credits,
    consume_data_processing,
    create_package_generation_order,
    fetch_package_banners,
    flush_pending_consumptions,
    get_package_generation_order,
    precheck_action,
)
from app.services.ai_analyzer import (
    _REQUIRED_AI_KEYS,
    compute_memo_content_hash,
    get_or_refresh_ai_result,
)
from app.services.analysis_result import (
    AnalysisResult,
    AnalyzedFile,
    ChildFolderSummary,
    ExtensionStat,
    FolderTreeNode,
    PriorityReviewFileCandidate,
    RecentModifiedFile,
)
from app.services.email_file_handler import process_email_files
from app.services.folder_scanner import scan_folder
from app.services.rag_package_builder import (
    EMBEDDING_BATCH_SIZE,
    RAG_TEXT_EXTRACTION_EXTENSIONS,
    RagPackageCancelled,
    build_and_save_rag_package,
    estimate_rag_package_cost,
    filter_files_by_selected_extensions,
    get_rag_package_candidate_files,
)
from app.services.report_writer import (
    save_analysis_result_as_json,
    save_analysis_result_as_word,
)
from app.ui.chatbot_dialog import ChatbotDialog
from app.ui.feedback_dialog import FeedbackDialog
from app.ui.memodialog import MemoDialog


_API_KEY_MASK = "●" * 8
_LICENSE_CODE_MASK = "●" * 8
_DAMAGED_LICENSE_CODE_MASK = "●" * 12
_API_KEY_ALREADY_SET_TOOLTIP = "API 키가 이미 설정되어 있습니다"
_LICENSE_PURCHASE_URL = "https://yourhr.co.kr"
_LICENSE_LOCK_DIALOG_TEXT = {
    "not_registered": (
        "라이선스 등록 필요",
        "아직 라이선스가 등록되지 않았습니다.<br>라이선스 키를 입력해 주세요.",
    ),
    "invalid_format": (
        "라이선스 코드 확인",
        "라이선스 코드 형식이 올바르지 않습니다.<br>담당 컨설턴트에게 문의해 주세요.",
    ),
    "no_internet": (
        "인터넷 연결 필요",
        "인터넷 연결이 필요합니다.<br>인터넷 연결 후 다시 실행해주세요.",
    ),
    "device_id_failed": (
        "PC 식별 실패",
        "PC 식별에 실패했습니다.<br>문제가 계속되면 문의해주세요.",
    ),
    "invalid_code": (
        "라이선스 등록",
        "입력하신 라이선스 코드를 찾을 수 없습니다.<br>코드를 다시 확인해 주세요.",
    ),
    "not_found": (
        "라이선스 등록",
        "입력하신 라이선스 코드를 찾을 수 없습니다.<br>코드를 다시 확인해 주세요.",
    ),
    "server_error": (
        "서버 연결 실패",
        "서버 연결에 실패했습니다.<br>인터넷 연결을 확인해주세요.",
    ),
    "other_device": (
        "라이선스 중복 사용",
        "이 라이선스는 이미 다른 PC에서 사용 중입니다.<br>재구매 또는 문의 부탁드립니다.",
    ),
    "license_terminated": (
        "라이선스 종료",
        (
            "라이선스가 종료되었습니다.<br>"
            "사유: 종료일이 지난 라이선스입니다.<br>"
            "담당 컨설턴트에게 문의해 주세요."
        ),
    ),
    "expired": (
        "라이선스 만료",
        (
            "라이선스가 만료되었습니다.<br>"
            "계속 사용하시려면 라이선스를 구매해주세요.<br><br>"
            f"<a href='{_LICENSE_PURCHASE_URL}'>구매 페이지 열기</a>"
        ),
    ),
}
_LICENSE_LOCK_LABEL_TEXT = {
    "not_registered": "아직 라이선스가 등록되지 않았습니다",
    "invalid_format": "라이선스 코드 형식이 올바르지 않습니다",
    "no_internet": "인터넷 연결이 필요합니다",
    "device_id_failed": "PC 식별에 실패했습니다",
    "invalid_code": "입력하신 라이선스 코드를 찾을 수 없습니다",
    "not_found": "입력하신 라이선스 코드를 찾을 수 없습니다",
    "server_error": "서버 연결에 실패했습니다",
    "other_device": "이 라이선스는 다른 PC에서 사용 중입니다",
    "license_terminated": "라이선스가 종료되었습니다",
    "expired": "라이선스가 만료되었습니다",
}
_API_KEY_GUIDE_PROMPT = """OpenAI API 키를 처음 발급받으려고 해. 나는 프로그래밍이나 개발
경험이 전혀 없는 완전 초보자야. 아래 내용을 하나도 빠짐없이,
아주 자세하고 쉽게 설명해줘.
1. OpenAI API 키를 발급받으려면 어느 웹사이트에 가야 하는지
   (정확한 주소와, 일반 ChatGPT 사이트와 다른 곳인지 여부)
2. 계정이 없다면 가입하는 방법 (이미 ChatGPT 계정이 있다면
   그걸 그대로 써도 되는지)
3. 결제 수단(신용카드 등)을 반드시 등록해야 하는지, 등록 안
   하면 어떻게 되는지
4. API 키를 실제로 생성하는 화면까지 가는 정확한 경로 (메뉴
   이름, 버튼 이름까지 구체적으로)
5. API 키 생성 시 나오는 이름 입력, 권한 설정 같은 옵션들을
   초보자는 어떻게 선택해야 하는지 (기본값을 써도 되는지)
6. 생성된 키를 어떻게 안전하게 보관해야 하는지 (키는 한 번만
   보여주고 다시 못 본다고 하는데 이게 무슨 뜻인지)
7. 이 키를 사용할 때 실제로 비용이 얼마나 드는지, 어떻게 하면
   예상치 못한 큰 비용이 나가는 걸 막을 수 있는지 (사용량 한도
   설정 방법 포함)
8. 키를 다른 사람에게 보여주거나 공개된 곳에 올리면 왜 위험한지,
   실수로 노출됐을 때 어떻게 해야 하는지
전문 용어를 쓸 때는 반드시 쉬운 말로 한 번 더 풀어서 설명해주고,
각 단계마다 '이렇게 하면 무슨 화면이 나온다'까지 알려줘. 지금
2026년 기준으로 최신 화면/절차로 알려줘."""
_API_KEY_GUIDE_BODY = f"""인수인계 프로그램에 사용하는 AI는 현재 가성비가 가장 뛰어난
GPT를 사용합니다. 아래 프롬프트를 사용하시는 AI에 요청하여,
OpenAI API를 발급 받고 입력하시기 바랍니다.

[사용하실 프롬프트]
{_API_KEY_GUIDE_PROMPT}"""
WINDOW_WIDTH = 633
WINDOW_MIN_HEIGHT = 727
OUTER_MARGIN = 6
MAIN_MARGIN = 20
MAIN_SECTION_SPACING = 20
MODE_SECTION_SPACING = 10
MODE_CARD_SPACING = 12
MODE_ACTION_SPACING = 12
ACTION_BUTTON_SPACING = 8
LIST_HEIGHT = 90
ACTION_SEPARATOR_HEIGHT = 1
FOOTER_SPACING = 25
_DOCUMENT_EXTENSIONS_FOR_PRIORITY_HINT = {
    ".doc",
    ".docx",
    ".xls",
    ".xlsx",
    ".pdf",
    ".ppt",
    ".pptx",
    ".hwp",
}
# A QListWidget with any stylesheet applied to it (even just border/border-radius)
# stops using the native checkbox indicator, and since no ::indicator style is
# defined here the native QListWidgetItem checkbox renders as a blank white
# block. Drawn as a ☐/☑ text symbol instead, same workaround as the folder
# tree in memodialog.py.
_LIST_ITEM_CHECK_STATE_ROLE = Qt.ItemDataRole.UserRole + 1
_UNCHECKED_SYMBOL = "☐ "
_CHECKED_SYMBOL = "☑ "


def _set_button_role(button: QPushButton, role: str) -> None:
    button.setProperty("buttonRole", role)


def _mask_license_code(license_code: str) -> str:
    parts = license_code.strip().rsplit("-", 2)
    if len(parts) != 3:
        return _DAMAGED_LICENSE_CODE_MASK

    company_code, validity_code, checksum = parts
    if not company_code or not validity_code or not checksum:
        return _DAMAGED_LICENSE_CODE_MASK

    return f"{company_code}-{_LICENSE_CODE_MASK}"


class _EqualWidthTabBar(QTabBar):
    disabledTabClicked = Signal(int)

    def tabSizeHint(self, index: int) -> QSize:
        size_hint = super().tabSizeHint(index)
        tab_count = max(self.count(), 1)
        return QSize(self.width() // tab_count, size_hint.height())

    def mousePressEvent(self, event) -> None:
        index = self.tabAt(event.position().toPoint())
        if index >= 0 and not self.isTabEnabled(index):
            self.disabledTabClicked.emit(index)
        super().mousePressEvent(event)


class _EqualWidthTabWidget(QTabWidget):
    # QTabWidget never stretches its tab bar past its own sizeHint, so on
    # resize the bar (and therefore each equal-width tab) would stay frozen
    # at its initial width. Force the bar to match the widget's full width.
    def resizeEvent(self, event) -> None:
        super().resizeEvent(event)
        self.tabBar().resize(self.width(), self.tabBar().height())


def _skip_confirmation_dialog(*_args, **_kwargs):
    return QMessageBox.StandardButton.Yes


class WorkflowProgressBar(QWidget):
    """Static 3-step guide for the main handover actions.

    "분석시작" no longer has its own step: [메모 작성 및 인수인계서 저장]
    now runs analysis automatically when needed, so that stage is folded
    into "메모작성".

    Intentionally never reflects live completion state (see
    MemoWorkflowProgressBar in memodialog.py for the same pattern) — this is
    a purely informational, always-neutral order-of-operations label.
    """

    STEP_LABELS = ("메모작성", "패키지생성", "물어보기")
    STEP_COLOR = QColor("#7C3AED")
    LABEL_COLOR = QColor("#6B7280")

    def __init__(self, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.setObjectName("workflowProgressBar")
        self.setMinimumHeight(62)
        self.setSizePolicy(
            self.sizePolicy().horizontalPolicy(),
            self.sizePolicy().verticalPolicy(),
        )

    def paintEvent(self, _event) -> None:
        painter = QPainter(self)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        width = max(self.width(), 1)
        step_count = len(self.STEP_LABELS)
        side_margin = max(42.0, min(76.0, width * 0.075))
        available_width = max(1.0, width - (side_margin * 2))
        centers = [
            QPointF(side_margin + available_width * index / (step_count - 1), 19.0)
            for index in range(step_count)
        ]

        painter.setPen(QPen(self.STEP_COLOR, 2.0, Qt.PenStyle.SolidLine))
        for index in range(step_count - 1):
            painter.drawLine(
                QPointF(centers[index].x() + 14, centers[index].y()),
                QPointF(centers[index + 1].x() - 14, centers[index + 1].y()),
            )

        symbol_font = QFont(self.font())
        symbol_font.setPixelSize(12)
        symbol_font.setBold(True)
        label_font = QFont(self.font())
        label_font.setPixelSize(11)

        for index, center in enumerate(centers):
            circle = QRectF(center.x() - 12, center.y() - 12, 24, 24)
            painter.setPen(QPen(self.STEP_COLOR, 2.0))
            painter.setBrush(QColor("#FFFFFF"))
            painter.drawEllipse(circle)

            painter.setFont(symbol_font)
            painter.setPen(self.STEP_COLOR)
            painter.drawText(circle, Qt.AlignmentFlag.AlignCenter, str(index + 1))

            painter.setFont(label_font)
            painter.setPen(self.LABEL_COLOR)
            label_rect = QRectF(center.x() - 58, 37, 116, 20)
            painter.drawText(
                label_rect,
                Qt.AlignmentFlag.AlignHCenter | Qt.AlignmentFlag.AlignTop,
                self.STEP_LABELS[index],
            )


class FileContentExtensionDialog(QDialog):
    EXTENSION_GROUPS = (
        ("워드 문서", (".docx",), 3, "워드(DOCX)"),
        ("파워포인트", (".pptx", ".ppt"), 3, "파워포인트(PPT/PPTX)"),
        ("한글", (".hwp", ".hwpx"), 8, "한글(HWP/HWPX)"),
        ("엑셀", (".xlsx", ".xls"), 12, "엑셀(XLS/XLSX)"),
        ("텍스트/마크다운", (".txt", ".md"), 1, "텍스트/마크다운"),
    )
    TIME_LEVEL_STYLES = {
        "낮음": "color: #475569; background: #F1F5F9; border: 1px solid #CBD5E1;",
        "중간": "color: #334155; background: #E2E8F0; border: 1px solid #CBD5E1;",
        "다소 높음": "color: #9A3412; background: #FFEDD5; border: 1px solid #FDBA74;",
        "매우 높음": "color: #B91C1C; background: #FEE2E2; border: 1px solid #FCA5A5;",
    }
    SIZE_LIMIT_OPTIONS = (
        ("안함", 0),
        ("1메가", 1 * 1024 * 1024),
        ("5메가", 5 * 1024 * 1024),
        ("10메가", 10 * 1024 * 1024),
        ("30메가", 30 * 1024 * 1024),
        ("제한없음", None),
    )
    LOCAL_SECONDS_PER_WEIGHTED_FILE = 88.6 / 2761
    ESTIMATED_BYTES_PER_CHUNK = 2048
    EMBEDDING_SECONDS_PER_BATCH = 5
    ESTIMATE_BUFFER_SECONDS = 5 * 60
    # Nominal probe size for the data-processing/check call used to read the
    # account's remaining quota (configuredFreeQuotaGb) — this field doesn't
    # vary with the requested size, so any small positive value works.
    REMAINING_QUOTA_PROBE_GB = 0.001
    SUMMARY_STYLE_NEUTRAL = (
        "QLabel { color: #0F172A; font-size: 13px; font-weight: 700; "
        "padding: 10px 12px; background: #EFF6FF; "
        "border: 1px solid #BFDBFE; border-radius: 6px; }"
    )
    SUMMARY_STYLE_WARNING = (
        "QLabel { color: #E57373; font-size: 13px; font-weight: 700; "
        "padding: 10px 12px; background: #FDEDED; "
        "border: 1px solid #F5C6CB; border-radius: 6px; }"
    )

    @staticmethod
    def processing_time_level(score: int) -> str:
        if score <= 250:
            return "낮음"
        if score <= 750:
            return "중간"
        if score <= 1500:
            return "다소 높음"
        return "매우 높음"

    def __init__(
        self,
        files: list[AnalyzedFile],
        parent: QWidget | None = None,
        *,
        extra_size_bytes: int = 0,
    ) -> None:
        super().__init__(parent)
        self._files = list(files)
        self._extra_size_bytes = max(0, int(extra_size_bytes))
        self._license_code = load_saved_license_code() or ""
        self._quota_state = "checking"  # "checking" | "ready" | "failed"
        self._remaining_gb: float | None = None
        self._quota_worker: PackageOrderWorker | None = None
        self.setWindowTitle("인수인계패키지로 만들 파일 선택")
        self.setModal(True)
        self.setMinimumSize(630, 300)
        counts = Counter(
            Path(file.file_name).suffix.lower()
            for file in files
            if Path(file.file_name).suffix.lower() in RAG_TEXT_EXTRACTION_EXTENSIONS
        )
        self._extension_limit_combos: list[tuple[QComboBox, tuple[str, ...]]] = []
        self._time_badges: dict[tuple[str, ...], QLabel] = {}

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok
            | QDialogButtonBox.StandardButton.Cancel,
            parent=self,
        )
        buttons.button(QDialogButtonBox.StandardButton.Ok).setText("확인")
        buttons.button(QDialogButtonBox.StandardButton.Cancel).setText("취소")
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        self.confirm_button = buttons.button(QDialogButtonBox.StandardButton.Ok)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(22, 20, 22, 18)
        layout.setSpacing(12)

        for label, extensions, weight, _guide_label in self.EXTENSION_GROUPS:
            count = sum(counts[extension] for extension in extensions)
            if count == 0:
                continue
            time_level = self.processing_time_level(count * weight)
            extension_text = ", ".join(extensions)
            group_label = QLabel(f"{label} ({extension_text}) — {count:,}개", self)
            group_label.setObjectName(f"extensionGroupLabel_{extensions[0][1:]}")
            group_label.setStyleSheet("QLabel { font-size: 12px; padding: 4px 2px; }")
            limit_combo = QComboBox(self)
            limit_combo.setObjectName(f"extensionLimitCombo_{extensions[0][1:]}")
            for option_label, size_bytes in self.SIZE_LIMIT_OPTIONS:
                limit_combo.addItem(option_label, size_bytes)
            limit_combo.setCurrentText("1메가")
            limit_combo.setMinimumWidth(105)
            limit_combo.setToolTip(
                "안함은 파일명만 포함하고, 나머지는 선택 용량 이하 파일의 내용을 포함합니다."
            )
            time_badge = QLabel(f"예상 시간: {time_level}", self)
            time_badge.setObjectName(f"processingTimeBadge_{extensions[0][1:]}")
            time_badge.setAlignment(Qt.AlignmentFlag.AlignCenter)
            time_badge.setStyleSheet(
                "QLabel { font-size: 11px; font-weight: 600; padding: 4px 8px; "
                "border-radius: 9px; "
                f"{self.TIME_LEVEL_STYLES[time_level]} }}"
            )
            row = QHBoxLayout()
            row.setContentsMargins(0, 0, 0, 0)
            row.setSpacing(0)
            row.addWidget(group_label)
            row.addSpacing(20)
            row.addWidget(time_badge)
            row.addStretch()
            row.addWidget(limit_combo)
            self._extension_limit_combos.append((limit_combo, extensions))
            self._time_badges[extensions] = time_badge
            limit_combo.currentIndexChanged.connect(self._update_estimated_time)
            layout.addLayout(row)

        self.estimated_time_label = QLabel(self)
        self.estimated_time_label.setObjectName("estimatedProcessingTimeLabel")
        self.estimated_time_label.setWordWrap(True)
        self.estimated_time_label.setStyleSheet(self.SUMMARY_STYLE_NEUTRAL)
        layout.addWidget(self.estimated_time_label)
        # The layout contributes its regular 12px spacing before this spacer,
        # producing an exact 25px visual gap above the purple benefit text.
        layout.addSpacing(13)
        self.package_benefit_label = QLabel(
            "챗봇을 사용하려면 '인수인계패키지 생성' 필수. "
            "다른 사람에게 물어보는걸 87% 줄이세요.",
            self,
        )
        self.package_benefit_label.setObjectName("packageBenefitLabel")
        self.package_benefit_label.setWordWrap(True)
        self.package_benefit_label.setStyleSheet(
            "QLabel { color: #7C3AED; font-size: 13px; font-weight: 600; "
            "padding: 0 12px; }"
        )
        layout.addWidget(self.package_benefit_label)
        self._start_quota_check()
        self._update_estimated_time()
        # The layout contributes its regular 12px spacing after this spacer,
        # producing an exact 25px visual gap below the purple benefit text.
        layout.addSpacing(13)
        layout.addWidget(buttons)

    def selected_extensions(self) -> set[str]:
        return {
            extension
            for combo, extensions in self._extension_limit_combos
            if combo.currentData() != 0
            for extension in extensions
        }

    def extension_size_limits(self) -> dict[str, int | None]:
        return {
            extension: combo.currentData()
            for combo, extensions in self._extension_limit_combos
            for extension in extensions
        }

    def content_target_files(self) -> list[AnalyzedFile]:
        content_files, _filename_only_files = filter_files_by_selected_extensions(
            self._files,
            self.selected_extensions(),
            extension_size_limits=self.extension_size_limits(),
        )
        return content_files

    @classmethod
    def _extension_weight(cls, file_name: str) -> int:
        extension = Path(file_name).suffix.lower()
        for _label, extensions, weight, _guide_label in cls.EXTENSION_GROUPS:
            if extension in extensions:
                return weight
        return 1

    @classmethod
    def _format_estimated_duration(cls, seconds: float) -> str:
        if seconds <= 0:
            return "약 0분 소요"
        if seconds <= 10 * 60:
            return "약 10분 소요"
        if seconds <= 20 * 60:
            return "약 15~20분 소요"
        if seconds <= 30 * 60:
            return "약 30분 소요"
        if seconds <= 60 * 60:
            return "약 1시간 소요"
        return f"약 {math.ceil(seconds / 3600)}시간 소요"

    def _estimated_duration_seconds(self, content_files: list[AnalyzedFile]) -> float:
        weighted_file_count = sum(
            self._extension_weight(file.file_name) for file in content_files
        )
        local_seconds = (
            weighted_file_count * self.LOCAL_SECONDS_PER_WEIGHTED_FILE
        )
        estimated_chunks = sum(
            max(1, math.ceil(file.size_bytes / self.ESTIMATED_BYTES_PER_CHUNK))
            for file in content_files
        )
        embedding_batches = (
            math.ceil(estimated_chunks / EMBEDDING_BATCH_SIZE)
            if estimated_chunks
            else 0
        )
        embedding_seconds = embedding_batches * self.EMBEDDING_SECONDS_PER_BATCH
        return (
            local_seconds + embedding_seconds + self.ESTIMATE_BUFFER_SECONDS
            if content_files
            else 0
        )

    def estimated_size_bytes(self) -> int:
        return self._extra_size_bytes + sum(
            file.size_bytes for file in self.content_target_files()
        )

    def _start_quota_check(self) -> None:
        if self._quota_worker is not None:
            return
        self._quota_state = "checking"
        if not self._license_code:
            self._quota_state = "failed"
            self._remaining_gb = None
            return
        worker = PackageOrderWorker(
            self._license_code, self.REMAINING_QUOTA_PROBE_GB, self
        )
        worker.completed.connect(self._handle_quota_check_completed)
        worker.finished.connect(worker.deleteLater)
        self._quota_worker = worker
        worker.start()

    @Slot(object)
    def _handle_quota_check_completed(self, result: object) -> None:
        self._quota_worker = None
        remaining_gb = None
        if isinstance(result, dict):
            try:
                remaining_gb = max(0.0, float(result["configuredFreeQuotaGb"]))
            except (KeyError, TypeError, ValueError):
                remaining_gb = None
        if remaining_gb is None:
            self._quota_state = "failed"
            self._remaining_gb = None
        else:
            self._quota_state = "ready"
            self._remaining_gb = remaining_gb
        self._update_estimated_time()

    @Slot()
    def _update_estimated_time(self, *_args) -> None:
        content_files = self.content_target_files()
        limits = self.extension_size_limits()
        for _combo, extensions in self._extension_limit_combos:
            matching_files = [
                file
                for file in content_files
                if Path(file.file_name).suffix.lower() in extensions
            ]
            badge = self._time_badges[extensions]
            if all(limits[extension] == 0 for extension in extensions):
                badge.setText("예상 시간: 해당 없음")
                badge.setStyleSheet(
                    "QLabel { color: #64748B; background: #F1F5F9; "
                    "font-size: 11px; font-weight: 600; padding: 4px 8px; border-radius: 9px; }"
                )
                continue
            score = sum(self._extension_weight(file.file_name) for file in matching_files)
            time_level = self.processing_time_level(score)
            badge.setText(f"예상 시간: {time_level}")
            badge.setStyleSheet(
                "QLabel { font-size: 11px; font-weight: 600; padding: 4px 8px; "
                "border-radius: 9px; "
                f"{self.TIME_LEVEL_STYLES[time_level]} }}"
            )
        duration_text = self._format_estimated_duration(
            self._estimated_duration_seconds(content_files)
        )
        size_gb = bytes_to_gb(self.estimated_size_bytes())

        if _args and self._quota_state == "failed" and self._quota_worker is None:
            # Give the user an easy retry: adjusting a combo re-attempts the
            # quota lookup instead of leaving it permanently stuck failed.
            # Gated on _args so this only fires for real combo-change calls
            # (which Qt invokes with the new index), not the internal
            # re-renders from __init__/_handle_quota_check_completed — those
            # must not loop straight back into another attempt.
            self._start_quota_check()

        if self._quota_state == "checking":
            quota_suffix = " | 자료 처리 확인 중..."
            style = self.SUMMARY_STYLE_NEUTRAL
            self.confirm_button.setEnabled(False)
        elif self._quota_state == "failed" or self._remaining_gb is None:
            quota_suffix = " | 자료 처리 확인 실패"
            style = self.SUMMARY_STYLE_NEUTRAL
            self.confirm_button.setEnabled(False)
        elif size_gb <= self._remaining_gb:
            quota_suffix = f" | 자료 처리 {self._remaining_gb:.2f}GB 가능"
            style = self.SUMMARY_STYLE_NEUTRAL
            self.confirm_button.setEnabled(True)
        else:
            quota_suffix = " | 처리 용량 결재 필요"
            style = self.SUMMARY_STYLE_WARNING
            self.confirm_button.setEnabled(False)

        self.estimated_time_label.setStyleSheet(style)
        self.estimated_time_label.setText(
            f"{len(content_files):,}개 파일(처리 용량 : {size_gb:.1f}GB) → "
            f"패키지 생성({duration_text}){quota_suffix}\n"
            "※ [문서 외 파일(동영상, 이미지, PDF, 캐드, 포토샵 등), 선택 용량 초과 파일]은 "
            "파일명/경로/수정일만 포함."
        )


class RagPackageProgressDialog(QDialog):
    def __init__(self, cancel_callback, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self._cancel_callback = cancel_callback
        self._cancel_requested = False
        self.banner_container: QWidget | None = None
        self.banner_labels: dict[str, QLabel] = {}
        self.banner_fetch_started = False

    def closeEvent(self, event) -> None:
        if self._cancel_requested:
            event.accept()
            return

        reply = QMessageBox.question(
            self,
            "인수인계패키지 생성",
            "패키지 생성을 취소하시겠습니까? 처리된 파일은 저장되어 다음 생성 때 이어서 진행됩니다.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No,
        )
        if reply != QMessageBox.StandardButton.Yes:
            event.ignore()
            return

        self._cancel_requested = True
        self._cancel_callback()
        event.accept()


class DataProcessingPaymentDialog(QDialog):
    retry_requested = Signal()
    cancel_requested = Signal()
    PORTAL_URL = "https://review.yourhr.co.kr/handover/portal"

    def __init__(self, shortfall_gb: float, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.setWindowTitle("자료 처리 결제 안내")
        self.setModal(False)
        self.setAttribute(Qt.WidgetAttribute.WA_DeleteOnClose, False)
        self.setFixedWidth(520)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(28, 26, 28, 24)
        layout.setSpacing(16)
        self.message_label = QLabel(self)
        self.message_label.setObjectName("dataProcessingPaymentMessage")
        self.message_label.setText(
            "자료 처리에 결제가 필요합니다. "
            f"필요 용량: {format_gb(shortfall_gb)}GB"
        )
        self.message_label.setWordWrap(True)
        self.message_label.setStyleSheet("font-size: 15px; font-weight: 700;")
        layout.addWidget(self.message_label)

        self.status_label = QLabel("", self)
        self.status_label.setObjectName("dataProcessingPaymentStatus")
        self.status_label.setWordWrap(True)
        self.status_label.setStyleSheet("color: #6B7280; font-size: 12px;")
        self.status_label.hide()
        layout.addWidget(self.status_label)

        row = QHBoxLayout()
        self.portal_button = QPushButton("설명서 페이지에서 결제하기", self)
        self.portal_button.setObjectName("openDataProcessingPortalButton")
        self.portal_button.clicked.connect(
            lambda: QDesktopServices.openUrl(QUrl(self.PORTAL_URL))
        )
        self.retry_button = QPushButton("이어서 진행", self)
        self.retry_button.setObjectName("retryDataProcessingButton")
        self.retry_button.clicked.connect(self.retry_requested.emit)
        self.cancel_button = QPushButton("취소", self)
        self.cancel_button.setObjectName("cancelDataProcessingButton")
        self.cancel_button.clicked.connect(self.cancel_requested.emit)
        row.addWidget(self.portal_button)
        row.addStretch()
        row.addWidget(self.retry_button)
        row.addWidget(self.cancel_button)
        layout.addLayout(row)

    def set_checking(self, checking: bool) -> None:
        self.retry_button.setEnabled(not checking)
        self.retry_button.setText("확인 중..." if checking else "이어서 진행")

    def show_not_confirmed(self) -> None:
        self.status_label.setText(
            "아직 결제가 확인되지 않았습니다. 결제 완료 후 다시 눌러주세요"
        )
        self.status_label.show()
        self.set_checking(False)

    def show_check_failed(self) -> None:
        self.status_label.setText("잔여량 확인에 실패했습니다. 잠시 후 다시 눌러주세요")
        self.status_label.show()
        self.set_checking(False)

    def closeEvent(self, event) -> None:
        event.ignore()
        self.cancel_requested.emit()


class RagPackageWorker(QThread):
    progress = Signal(str, int, int)
    succeeded = Signal(str, int, int, int)
    failed = Signal(str)
    canceled = Signal()

    def __init__(
        self,
        analysis_result: AnalysisResult,
        folder_paths: list[str],
        api_key: str,
        output_path: str,
        parsed_emails: list[dict],
        kakao_file_paths: list[str],
        extension_size_limits: dict[str, int | None],
        parent: QWidget | None = None,
    ) -> None:
        super().__init__(parent)
        self.analysis_result = analysis_result
        self.folder_paths = folder_paths
        self.api_key = api_key
        self.output_path = output_path
        self.parsed_emails = parsed_emails
        self.kakao_file_paths = kakao_file_paths
        self.extension_size_limits = extension_size_limits
        self._cancel_requested = False

    def request_cancel(self) -> None:
        self._cancel_requested = True

    def _is_cancel_requested(self) -> bool:
        return self._cancel_requested

    def _cleanup_partial_output(self) -> None:
        zip_path = Path(self.output_path)
        if zip_path.suffix.casefold() != ".zip":
            zip_path = zip_path.with_suffix(".zip")
        try:
            if zip_path.exists():
                zip_path.unlink()
        except OSError:
            pass

    def run(self) -> None:
        try:
            result = build_and_save_rag_package(
                self.analysis_result,
                self.folder_paths,
                self.api_key,
                self.output_path,
                parsed_emails=self.parsed_emails,
                kakao_file_paths=self.kakao_file_paths,
                selected_extensions=set(self.extension_size_limits),
                extension_size_limits=self.extension_size_limits,
                progress_callback=self.progress.emit,
                cancel_check=self._is_cancel_requested,
            )
        except RagPackageCancelled:
            self._cleanup_partial_output()
            self.canceled.emit()
            return
        except Exception as exc:
            self.failed.emit(str(exc))
            return
        if self._is_cancel_requested():
            self._cleanup_partial_output()
            self.canceled.emit()
            return
        failed_count = int(result.get("embedding_failed_chunk_count", 0))
        self.succeeded.emit(
            str(result.get("saved_path") or self.output_path),
            failed_count,
            int(result.get("embedding_tokens", 0)),
            int(result.get("timed_out_file_count", 0)),
        )


class CostEstimationWorker(QThread):
    succeeded = Signal(dict, list)
    failed = Signal(str)
    canceled = Signal()
    progress = Signal(str, int, int)

    def __init__(
        self,
        analysis_result: AnalysisResult,
        folder_paths: list[str],
        email_file_paths: list[str],
        kakao_file_paths: list[str],
        extension_size_limits: dict[str, int | None],
        parent: QWidget | None = None,
    ) -> None:
        super().__init__(parent)
        self.analysis_result = analysis_result
        self.folder_paths = folder_paths
        self.email_file_paths = email_file_paths
        self.kakao_file_paths = kakao_file_paths
        self.extension_size_limits = extension_size_limits
        self._cancel_requested = False

    def request_cancel(self) -> None:
        self._cancel_requested = True

    def _is_cancel_requested(self) -> bool:
        return self._cancel_requested

    def run(self) -> None:
        try:
            parsed_emails, _ = (
                process_email_files(self.email_file_paths)
                if self.email_file_paths
                else ([], 0)
            )
            estimate = estimate_rag_package_cost(
                self.analysis_result,
                self.folder_paths,
                parsed_emails,
                self.kakao_file_paths,
                cancel_check=self._is_cancel_requested,
                progress_callback=self.progress.emit,
                selected_extensions=set(self.extension_size_limits),
                extension_size_limits=self.extension_size_limits,
            )
        except RagPackageCancelled:
            self.canceled.emit()
            return
        except Exception as exc:
            self.failed.emit(str(exc))
            return
        if self._is_cancel_requested():
            self.canceled.emit()
            return
        self.succeeded.emit(estimate, parsed_emails)


class CandidateScanWorker(QThread):
    """Runs the folder scan + content-hash dedup pass off the UI thread.

    This is the same work _continue_create_rag_package used to run inline
    before showing FileContentExtensionDialog — scanning the folder tree and
    then SHA-256-hashing every candidate file's full content for dedup. On a
    large folder that can take from seconds to minutes, so it now runs here
    instead of blocking the GUI thread before the dialog can even appear.
    """

    succeeded = Signal(object, list)
    failed = Signal(str)
    canceled = Signal()

    def __init__(
        self,
        folder_paths: list[str],
        analysis_mode: str,
        parent: QWidget | None = None,
    ) -> None:
        super().__init__(parent)
        self.folder_paths = folder_paths
        self.analysis_mode = analysis_mode
        self._cancel_requested = False

    def request_cancel(self) -> None:
        self._cancel_requested = True

    def _is_cancel_requested(self) -> bool:
        return self._cancel_requested

    def run(self) -> None:
        try:
            if self.folder_paths:
                result = AnalysisWorker(
                    self.folder_paths, self.analysis_mode
                )._build_merged_analysis_result(self.folder_paths)
            else:
                result = AnalysisResult(
                    root_folder_path="",
                    total_folder_count=0,
                    total_file_count=0,
                    total_size_bytes=0,
                    modified_within_7_days_count=0,
                    modified_within_30_days_count=0,
                    modified_within_90_days_count=0,
                    error_count=0,
                    child_folder_summaries=[],
                )
            if self._is_cancel_requested():
                self.canceled.emit()
                return
            candidate_files = get_rag_package_candidate_files(
                result,
                self.folder_paths,
                cancel_check=self._is_cancel_requested,
            )
        except RagPackageCancelled:
            self.canceled.emit()
            return
        except Exception as exc:
            self.failed.emit(str(exc))
            return
        if self._is_cancel_requested():
            self.canceled.emit()
            return
        self.succeeded.emit(result, candidate_files)


class CandidateScanProgressDialog(QDialog):
    """Small cancelable loading dialog shown while CandidateScanWorker runs."""

    cancel_requested = Signal()

    def __init__(self, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.setWindowTitle("인수인계패키지 생성")
        self.setModal(True)
        self.setFixedSize(420, 150)
        self._resolved = False

        self._base_text = "선택한 폴더의 파일을 확인하고 있습니다"
        self._frame = 0

        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 24, 24, 20)
        layout.addStretch()

        self.status_label = QLabel(self._base_text, self)
        self.status_label.setObjectName("candidateScanStatusLabel")
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.status_label.setWordWrap(True)
        self.status_label.setStyleSheet("font-size: 14px; font-weight: 600;")
        layout.addWidget(self.status_label)

        self.hint_label = QLabel(
            "대용량 폴더는 중복 파일 확인에 다소 시간이 걸릴 수 있습니다.",
            self,
        )
        self.hint_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.hint_label.setWordWrap(True)
        self.hint_label.setStyleSheet("color: #6B7280; font-size: 11px;")
        layout.addWidget(self.hint_label)
        layout.addStretch()

        self.cancel_button = QPushButton("취소", self)
        self.cancel_button.setObjectName("candidateScanCancelButton")
        self.cancel_button.setAutoDefault(False)
        self.cancel_button.clicked.connect(self._handle_cancel_clicked)
        button_row = QHBoxLayout()
        button_row.addStretch()
        button_row.addWidget(self.cancel_button)
        layout.addLayout(button_row)

        self._animation_timer = QTimer(self)
        self._animation_timer.setInterval(350)
        self._animation_timer.timeout.connect(self._advance_animation)
        self._animation_timer.start()
        self._advance_animation()

    def _advance_animation(self) -> None:
        dots = ("...", "..", ".")[self._frame]
        self.status_label.setText(f"{self._base_text}{dots}")
        self._frame = (self._frame + 1) % 3

    def _handle_cancel_clicked(self) -> None:
        self.cancel_button.setEnabled(False)
        self.status_label.setText("취소하는 중...")
        self.resolve_canceled()

    def resolve_succeeded(self) -> None:
        if self._resolved:
            return
        self._resolved = True
        self._animation_timer.stop()
        self.accept()

    def resolve_failed(self) -> None:
        if self._resolved:
            return
        self._resolved = True
        self._animation_timer.stop()
        self.reject()

    def resolve_canceled(self) -> None:
        if self._resolved:
            return
        self._resolved = True
        self._animation_timer.stop()
        self.cancel_requested.emit()
        self.reject()

    def closeEvent(self, event) -> None:
        self.resolve_canceled()
        event.accept()


class CreditPrecheckWorker(QThread):
    completed = Signal(object)

    def __init__(self, license_code: str, action_type: str, parent=None) -> None:
        super().__init__(parent)
        self.license_code = license_code
        self.action_type = action_type

    def run(self) -> None:
        try:
            result = precheck_action(self.license_code, self.action_type)
        except Exception:
            result = None
        self.completed.emit(result)


class PackageBannerWorker(QThread):
    completed = Signal(object)

    def run(self) -> None:
        try:
            result = fetch_package_banners()
        except Exception:
            result = None
        self.completed.emit(result)


class CreditBalanceWorker(QThread):
    completed = Signal(str, object)

    def __init__(self, license_code: str, *, flush_pending: bool = False, parent=None) -> None:
        super().__init__(parent)
        self.license_code = license_code
        self.flush_pending = flush_pending

    def run(self) -> None:
        try:
            if self.flush_pending:
                flush_pending_consumptions()
            balance = check_balance(self.license_code)
        except Exception:
            balance = None
        self.completed.emit(self.license_code, balance)


class CreditFinalizeWorker(QThread):
    completed = Signal(str, object)

    def __init__(self, license_code: str, action_type: str, usage: dict, parent=None) -> None:
        super().__init__(parent)
        self.license_code = license_code
        self.action_type = action_type
        self.usage = usage

    def run(self) -> None:
        try:
            consume_credits(self.license_code, self.action_type, **self.usage)
            flush_pending_consumptions()
            balance = check_balance(self.license_code)
        except Exception:
            balance = None
        self.completed.emit(self.license_code, balance)


class PackageOrderWorker(QThread):
    completed = Signal(object)

    def __init__(self, license_code: str, size_gb: float, parent=None) -> None:
        super().__init__(parent)
        self.license_code = license_code
        self.size_gb = size_gb

    def run(self) -> None:
        self.completed.emit(
            create_package_generation_order(self.license_code, self.size_gb)
        )


class DataProcessingConsumeWorker(QThread):
    def __init__(self, license_code: str, actual_gb: float, parent=None) -> None:
        super().__init__(parent)
        self.license_code = license_code
        self.actual_gb = actual_gb

    def run(self) -> None:
        consume_data_processing(self.license_code, self.actual_gb)


class PackagePaymentPollWorker(QThread):
    completed = Signal(object)

    def __init__(self, order_id: str, timeout_seconds: int = 30 * 60, parent=None) -> None:
        super().__init__(parent)
        self.order_id = order_id
        self.timeout_seconds = timeout_seconds

    def run(self) -> None:
        deadline = time.monotonic() + self.timeout_seconds
        while not self.isInterruptionRequested() and time.monotonic() < deadline:
            result = get_package_generation_order(self.order_id)
            if result and result.get("status") in {"completed", "expired", "cancelled"}:
                self.completed.emit(result)
                return
            for _ in range(20):
                if self.isInterruptionRequested():
                    return
                self.msleep(250)
        if not self.isInterruptionRequested():
            self.completed.emit({"status": "timeout"})


class ReportAiWorker(QThread):
    succeeded = Signal(str, object)
    denied = Signal()
    failed = Signal(str)

    def __init__(self, pending_memos, all_files, root_folder_path, parsed_emails, license_code, parent=None):
        super().__init__(parent)
        self.pending_memos = pending_memos
        self.all_files = all_files
        self.root_folder_path = root_folder_path
        self.parsed_emails = parsed_emails
        self.license_code = license_code

    def run(self) -> None:
        try:
            for _memo in self.pending_memos:
                precheck = precheck_action(self.license_code, "report")
                if precheck is not None and precheck.get("allowed") is False:
                    self.denied.emit()
                    return
            for memo in self.pending_memos:
                ai_result = get_or_refresh_ai_result(
                    memo, self.all_files, self.root_folder_path, self.parsed_emails
                )
                if ai_result is not None:
                    usage = ai_result.get("_usage", {})
                    consume_credits(
                        self.license_code,
                        "report",
                        prompt_tokens=usage.get("prompt_tokens", 0),
                        completion_tokens=usage.get("completion_tokens", 0),
                    )
                    flush_pending_consumptions()
            self.succeeded.emit(self.license_code, check_balance(self.license_code))
        except Exception as exc:
            self.failed.emit(str(exc))


class AnalysisWorker(QThread):
    succeeded = Signal(object)
    failed = Signal(str)

    def __init__(
        self,
        folder_paths: list[str],
        analysis_mode: str,
        parent: QWidget | None = None,
    ) -> None:
        super().__init__(parent)
        self.folder_paths = folder_paths
        self.analysis_mode = analysis_mode

    def run(self) -> None:
        try:
            result = replace(
                self._build_merged_analysis_result(self.folder_paths),
                analysismode=self.analysis_mode,
            )
        except Exception as exc:
            self.failed.emit(str(exc))
            return
        self.succeeded.emit(result)

    def _build_merged_analysis_result(self, folder_paths: list[str]) -> AnalysisResult:
        if len(folder_paths) == 1:
            sub_result = scan_folder(folder_paths[0])
            namespace = Path(folder_paths[0]).name or folder_paths[0]
            return self._wrap_single_folder_analysis_result(
                sub_result,
                folder_paths[0],
                namespace,
            )

        used_namespaces: set[str] = set()
        child_folder_summaries: list[ChildFolderSummary] = []
        folder_tree: list[FolderTreeNode] = []
        all_files: list[AnalyzedFile] = []
        total_folder_count = 0
        total_file_count = 0
        total_size_bytes = 0
        modified_within_7_days_count = 0
        modified_within_30_days_count = 0
        modified_within_90_days_count = 0
        error_count = 0

        for folder_path in folder_paths:
            sub_result = scan_folder(folder_path)
            namespace = self._make_unique_namespace(
                Path(folder_path).name or folder_path,
                used_namespaces,
            )

            total_folder_count += sub_result.total_folder_count
            total_file_count += sub_result.total_file_count
            total_size_bytes += sub_result.total_size_bytes
            modified_within_7_days_count += sub_result.modified_within_7_days_count
            modified_within_30_days_count += sub_result.modified_within_30_days_count
            modified_within_90_days_count += sub_result.modified_within_90_days_count
            error_count += sub_result.error_count

            child_folder_summaries.append(
                self._build_namespace_child_summary(namespace, sub_result)
            )
            folder_tree.append(
                FolderTreeNode(
                    name=namespace,
                    relative_path=namespace,
                    children=self._reprefix_folder_tree(
                        sub_result.folder_tree, namespace
                    ),
                )
            )
            all_files.extend(
                replace(file, relative_path=f"{namespace}/{file.relative_path}")
                for file in sub_result.all_files
            )

        return AnalysisResult(
            root_folder_path="; ".join(folder_paths),
            total_folder_count=total_folder_count,
            total_file_count=total_file_count,
            total_size_bytes=total_size_bytes,
            modified_within_7_days_count=modified_within_7_days_count,
            modified_within_30_days_count=modified_within_30_days_count,
            modified_within_90_days_count=modified_within_90_days_count,
            error_count=error_count,
            child_folder_summaries=child_folder_summaries,
            folder_tree=folder_tree,
            all_files=all_files,
        )

    def _wrap_single_folder_analysis_result(
        self,
        sub_result: AnalysisResult,
        folder_path: str,
        namespace: str,
    ) -> AnalysisResult:
        return AnalysisResult(
            root_folder_path=folder_path,
            total_folder_count=sub_result.total_folder_count,
            total_file_count=sub_result.total_file_count,
            total_size_bytes=sub_result.total_size_bytes,
            modified_within_7_days_count=sub_result.modified_within_7_days_count,
            modified_within_30_days_count=sub_result.modified_within_30_days_count,
            modified_within_90_days_count=sub_result.modified_within_90_days_count,
            error_count=sub_result.error_count,
            child_folder_summaries=[
                self._build_namespace_child_summary(namespace, sub_result)
            ],
            folder_tree=[
                FolderTreeNode(
                    name=namespace,
                    relative_path=namespace,
                    children=self._reprefix_folder_tree(
                        sub_result.folder_tree, namespace
                    ),
                )
            ],
            memos=sub_result.memos,
            all_files=[
                replace(file, relative_path=f"{namespace}/{file.relative_path}")
                for file in sub_result.all_files
            ],
            analysismode=sub_result.analysismode,
        )

    def _make_unique_namespace(self, base_name: str, used_namespaces: set[str]) -> str:
        candidate = base_name
        suffix = 2
        while candidate in used_namespaces:
            candidate = f"{base_name} ({suffix})"
            suffix += 1
        used_namespaces.add(candidate)
        return candidate

    def _reprefix_folder_tree(
        self,
        nodes: list[FolderTreeNode],
        namespace: str,
    ) -> list[FolderTreeNode]:
        return [
            FolderTreeNode(
                name=node.name,
                relative_path=f"{namespace}/{node.relative_path}",
                children=self._reprefix_folder_tree(node.children, namespace),
            )
            for node in nodes
        ]

    def _build_namespace_child_summary(
        self,
        namespace: str,
        sub_result: AnalysisResult,
    ) -> ChildFolderSummary:
        files = sub_result.all_files

        recent_modified_files = [
            RecentModifiedFile(
                file_name=file.file_name,
                relative_path=f"{namespace}/{file.relative_path}",
                modified_at=file.modified_at,
            )
            for file in sorted(
                files, key=lambda f: f.modified_timestamp, reverse=True
            )[:10]
        ]

        extension_counts: dict[str, int] = {}
        for file in files:
            extension = Path(file.file_name).suffix.lower() or "[no extension]"
            extension_counts[extension] = extension_counts.get(extension, 0) + 1
        extension_stats = [
            ExtensionStat(extension=extension, file_count=file_count)
            for extension, file_count in sorted(
                extension_counts.items(), key=lambda item: (-item[1], item[0])
            )[:10]
        ]

        candidate_pool = [
            file
            for file in files
            if file.size_bytes > 0
            and not file.is_hidden_or_system
            and not file.file_name.startswith("~$")
        ]
        candidate_pool.sort(
            key=lambda f: (
                Path(f.file_name).suffix.lower()
                in _DOCUMENT_EXTENSIONS_FOR_PRIORITY_HINT,
                f.modified_timestamp,
            ),
            reverse=True,
        )
        priority_review_file_candidates = [
            PriorityReviewFileCandidate(
                file_name=file.file_name,
                relative_path=f"{namespace}/{file.relative_path}",
                modified_at=file.modified_at,
                size_bytes=file.size_bytes,
            )
            for file in candidate_pool[:5]
        ]

        return ChildFolderSummary(
            folder_name=namespace,
            relative_path=namespace,
            total_folder_count=sub_result.total_folder_count,
            total_file_count=sub_result.total_file_count,
            total_size_bytes=sub_result.total_size_bytes,
            modified_within_30_days_count=sub_result.modified_within_30_days_count,
            recent_modified_files=recent_modified_files,
            extension_stats=extension_stats,
            priority_review_file_candidates=priority_review_file_candidates,
        )


class MainWindow(QMainWindow):
    def __init__(self) -> None:
        super().__init__()

        self.setWindowTitle("인수인계 프로그램")
        self.resize(WINDOW_WIDTH, WINDOW_MIN_HEIGHT)
        self.setMinimumWidth(WINDOW_WIDTH)

        self.current_analysis_result: AnalysisResult | None = None
        self.current_analysis_analyzed_at: datetime | None = None
        # (folder paths, email paths, kakao paths) signature of the selection
        # that current_analysis_result was actually produced from — compared
        # against the live selection so [메모 작성 및 인수인계서 저장] can skip
        # re-analysis when nothing changed. Set only once analysis succeeds
        # (see _handle_analysis_succeeded), paired with _pending_..., which is
        # captured right before the worker starts so a failed/in-flight run
        # never leaves the two out of sync with current_analysis_result.
        self._analyzed_selection_signature: (
            tuple[tuple[str, ...], tuple[str, ...], tuple[str, ...]] | None
        ) = None
        self._pending_analysis_selection_signature: (
            tuple[tuple[str, ...], tuple[str, ...], tuple[str, ...]] | None
        ) = None
        self._chatbot_dialog: ChatbotDialog | None = None
        self._memo_dialog: MemoDialog | None = None
        self._last_saved_report_fingerprint: str | None = None
        self._last_saved_word_path: str | None = None
        self._rag_package_context: dict[str, object] | None = None
        self._cost_estimation_worker: CostEstimationWorker | None = None
        self._rag_package_worker: RagPackageWorker | None = None
        self._rag_package_progress_box: QDialog | None = None
        self._rag_package_progress_label: QLabel | None = None
        self._rag_package_progress_timer: QTimer | None = None
        self._analysis_worker: AnalysisWorker | None = None
        self._analysis_progress_box: QMessageBox | None = None
        self._analysis_progress_timer: QTimer | None = None
        self._credit_workers: set[QThread] = set()
        self._package_banner_workers: set[PackageBannerWorker] = set()
        self._active_extension_dialog: FileContentExtensionDialog | None = None
        self._candidate_scan_worker: CandidateScanWorker | None = None
        self._package_order_worker: PackageOrderWorker | None = None
        self._data_processing_payment_dialog: DataProcessingPaymentDialog | None = None
        self._data_processing_consume_worker: DataProcessingConsumeWorker | None = None
        self._package_payment_worker: PackagePaymentPollWorker | None = None
        self._report_ai_worker: ReportAiWorker | None = None
        self._report_ai_progress_box: QMessageBox | None = None
        self._pending_word_save: tuple[str, bool] | None = None
        self._credit_refresh_timer: QTimer | None = None

        self.app_title_label = QLabel(
            '<span style="font-size:22px; font-weight:bold; color:#1A1A1A;">'
            "인수인계 쉽게, 정확하게</span> "
            '<span style="font-size:16px; font-weight:bold; color:#1A1A1A;">'
            "10분 완성</span>"
        )
        self.brand_label = QLabel(
            "<span style='font-size:10px; font-weight:bold; color:#1A1A1A;'>yourHR</span> "
            "<span style='font-size:10px; font-weight:bold; color:#7C3AED;'>대표님의 인사담당자</span>"
        )
        self.brand_label.setAlignment(Qt.AlignmentFlag.AlignRight)
        self.brand_label.setMinimumWidth(1)
        self.copyright_label = QLabel(
            "© 2026 yourHR 대표님의 인사담당자. All Rights Reserved."
        )
        self.copyright_label.setObjectName("copyrightLabel")
        self.copyright_label.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.folder_path_input = QLineEdit()
        self.folder_path_input.setObjectName("folderPathInput")
        self.folder_path_input.setPlaceholderText("분석 대상 폴더 경로")

        self.select_folder_button = QPushButton("폴더 선택")
        _set_button_role(self.select_folder_button, "secondary")

        self.remove_selected_folders_button = QPushButton("선택 삭제")
        _set_button_role(self.remove_selected_folders_button, "secondary")
        self.remove_all_folders_button = QPushButton("전체 삭제")
        _set_button_role(self.remove_all_folders_button, "secondary")

        self.folder_list_widget = QListWidget()
        self.folder_list_widget.setObjectName("mainFolderList")
        self.folder_list_widget.setFixedHeight(LIST_HEIGHT)
        self.folder_empty_label = self._create_empty_state_label(
            "선택된 폴더가 없습니다"
        )

        self.edit_memo_button = QPushButton("메모 작성 및 인수인계서 저장")
        _set_button_role(self.edit_memo_button, "primary")
        self.edit_memo_button.setEnabled(True)
        self.create_rag_package_button = QPushButton("인수인계패키지 생성")
        _set_button_role(self.create_rag_package_button, "secondary")
        self.create_rag_package_button.setEnabled(True)
        self.chatbot_button = QPushButton("물어보기")
        _set_button_role(self.chatbot_button, "secondary")
        self.feedback_button = QPushButton("💬 의견 보내기")
        self.feedback_button.setObjectName("feedbackButton")
        self.feedback_button.setFixedHeight(28)
        self.feedback_button.setMaximumWidth(110)
        self.feedback_button.setStyleSheet(
            "QPushButton { color: #64748B; background: #F8FAFC; "
            "border: 1px solid #E2E8F0; border-radius: 6px; "
            "font-size: 11px; padding: 4px 9px; } "
            "QPushButton:hover { color: #475569; background: #F1F5F9; "
            "border-color: #CBD5E1; }"
        )
        self.workflow_progress = WorkflowProgressBar(self)
        self.save_json_button = QPushButton("JSON 저장")
        _set_button_role(self.save_json_button, "secondary")

        self.selected_analysis_mode = "basic"
        self.mode_cards: dict[str, tuple[QFrame, QLabel]] = {}
        self.license_activated, self._license_lock_reason = (
            self._evaluate_startup_license_state()
        )
        self.license_unlock_button = QPushButton("라이선스 등록")
        _set_button_role(self.license_unlock_button, "secondary")
        self.api_key = load_api_key()
        self.api_key_button = QPushButton("GPT API 키")
        _set_button_role(self.api_key_button, "secondary")
        self._refresh_api_key_button_state()

        self.analysis_mode_notice = QLabel("")
        self.analysis_mode_notice.setObjectName("analysisModeNotice")
        self.license_status_label = QLabel(
            _LICENSE_LOCK_LABEL_TEXT.get(self._license_lock_reason, "라이선스가 만료되었습니다")
        )
        self.license_status_label.setObjectName("licenseStatusLabel")
        self.license_status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.license_status_label.setVisible(not self.license_activated)
        self.credit_balance_banner = QLabel("")
        self.credit_balance_banner.setObjectName("creditBalanceBanner")
        self.credit_balance_banner.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.credit_balance_banner.setWordWrap(True)
        self.credit_balance_banner.setStyleSheet(
            "QLabel { color: #92400e; background: #fef3c7; border: 1px solid #f59e0b; "
            "border-radius: 6px; padding: 6px 10px; }"
        )
        self.credit_balance_banner.hide()

        self.email_format_notice_label = QLabel(
            "지원 형식 : .eml, .msg, .zip (zip 안에 eml/msg 포함 가능) ·"
            " .pst는 지원 X"
        )
        self.email_format_notice_label.setObjectName("emailFormatNotice")
        self.email_format_notice_label.setWordWrap(True)

        self.add_email_files_button = QPushButton("메일 파일 추가")
        _set_button_role(self.add_email_files_button, "secondary")

        self.remove_selected_email_files_button = QPushButton("선택 삭제")
        _set_button_role(self.remove_selected_email_files_button, "secondary")
        self.remove_all_email_files_button = QPushButton("전체 삭제")
        _set_button_role(self.remove_all_email_files_button, "secondary")

        self.email_file_list_widget = QListWidget()
        self.email_file_list_widget.setObjectName("mainEmailList")
        self.email_file_list_widget.setFixedHeight(LIST_HEIGHT)
        self.email_empty_label = self._create_empty_state_label(
            "선택된 메일이 없습니다"
        )

        self.kakao_format_notice_label = QLabel(
            "지원 형식 : 카카오톡 등 각종 메신져의 대화 백업 파일(txt 파일)"
        )
        self.kakao_format_notice_label.setObjectName("kakaoFormatNotice")
        self.kakao_format_notice_label.setWordWrap(True)

        self.add_kakao_files_button = QPushButton("메신저(카톡) 파일 추가")
        _set_button_role(self.add_kakao_files_button, "secondary")

        self.remove_selected_kakao_files_button = QPushButton("선택 삭제")
        _set_button_role(self.remove_selected_kakao_files_button, "secondary")
        self.remove_all_kakao_files_button = QPushButton("전체 삭제")
        _set_button_role(self.remove_all_kakao_files_button, "secondary")

        self.kakao_file_list_widget = QListWidget()
        self.kakao_file_list_widget.setObjectName("mainKakaoList")
        self.kakao_file_list_widget.setFixedHeight(LIST_HEIGHT)
        self.kakao_empty_label = self._create_empty_state_label(
            "선택된 메신저(카톡) 파일이 없습니다"
        )

        self.result_preview = QTextEdit()
        self.result_preview.setPlaceholderText("결과 미리보기")
        self.result_preview.hide()
        self.result_preview_label = QLabel("결과 미리보기")
        self.result_preview_label.hide()
        self.save_json_button.hide()

        self._build_layout()
        self._connect_signals()
        self._apply_license_lock_state(show_warning=not self.license_activated)

    def _create_mode_card(
        self,
        mode_key: str,
        title: str,
        description: str,
    ) -> QFrame:
        card = QFrame()
        card.setObjectName("modeCard")
        card.setProperty("selected", "false")
        card.setCursor(Qt.CursorShape.PointingHandCursor)
        card.mousePressEvent = (
            lambda _event, key=mode_key: self._select_analysis_mode(key)
        )

        card_layout = QHBoxLayout(card)
        card_layout.setContentsMargins(12, 10, 12, 10)
        card_layout.setSpacing(10)

        symbol_label = QLabel("●" if mode_key == self.selected_analysis_mode else "○")
        symbol_label.setObjectName("modeCardSymbol")
        symbol_label.setProperty("selected", "false")
        card_layout.addWidget(symbol_label, 0, Qt.AlignmentFlag.AlignVCenter)

        text_layout = QVBoxLayout()
        text_layout.setSpacing(2)
        title_label = QLabel(title)
        title_label.setObjectName("modeCardTitle")
        description_label = QLabel(description)
        description_label.setObjectName("modeCardDescription")
        description_label.setWordWrap(False)
        text_layout.addWidget(title_label)
        text_layout.addWidget(description_label)
        card_layout.addLayout(text_layout)
        card_layout.addStretch()

        self.mode_cards[mode_key] = (card, symbol_label)
        return card

    def _create_empty_state_label(self, text: str) -> QLabel:
        label = QLabel(text)
        label.setObjectName("emptyStateLabel")
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        return label

    def _create_list_stack(
        self,
        list_widget: QListWidget,
        empty_label: QLabel,
    ) -> QWidget:
        container = QWidget()
        container.setFixedHeight(LIST_HEIGHT)
        stack_layout = QStackedLayout(container)
        stack_layout.setContentsMargins(0, 0, 0, 0)
        stack_layout.addWidget(empty_label)
        stack_layout.addWidget(list_widget)
        return container

    def _update_empty_state_labels(self) -> None:
        for list_widget, empty_label in (
            (self.folder_list_widget, self.folder_empty_label),
            (self.email_file_list_widget, self.email_empty_label),
            (self.kakao_file_list_widget, self.kakao_empty_label),
        ):
            stack_layout = list_widget.parentWidget().layout()
            if not isinstance(stack_layout, QStackedLayout):
                continue
            stack_layout.setCurrentWidget(
                empty_label if list_widget.count() == 0 else list_widget
            )

    def _select_analysis_mode(self, mode_key: str) -> None:
        if not self.license_activated:
            self._show_license_lock_warning()
            return

        previous_mode = self.selected_analysis_mode
        if previous_mode == mode_key:
            return
        if self._has_analysis_targets_or_memo_content():
            reply = QMessageBox.question(
                self,
                "분석 모드 변경",
                "모드를 변경하면 선택한 내용과 작성한 메모가 모두 초기화 됩니다.",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No,
            )
            if reply != QMessageBox.StandardButton.Yes:
                return

        self._reset_for_analysis_mode_change()
        self.selected_analysis_mode = mode_key
        self._update_mode_card_styles()
        self._update_analysis_mode_notice()
        self._update_analysis_target_tabs_enabled()

    def _has_analysis_targets_or_memo_content(self) -> bool:
        return bool(
            self._get_selected_folder_paths()
            or self._get_selected_email_file_paths()
            or self._get_selected_kakao_file_paths()
            or self._has_memo_content_or_links()
        )

    def _has_memo_content_or_links(self) -> bool:
        result = self.current_analysis_result
        if result is None:
            return False
        return any(
            memo.title.strip()
            or memo.content.strip()
            or memo.linked_folders
            or memo.linked_files
            or memo.linked_emails
            or memo.linked_kakao_files
            for memo in result.memos
        )

    def _confirm_clear_memos_for_analysis_restart(self) -> bool:
        if not self._has_memo_content_or_links():
            return True
        reply = QMessageBox.question(
            self,
            "분석 시작",
            "분석시작을 다시 하면 메모작성팝업의 내용이 모두 삭제됩니다.",
            QMessageBox.StandardButton.Ok | QMessageBox.StandardButton.Cancel,
            QMessageBox.StandardButton.Cancel,
        )
        return reply == QMessageBox.StandardButton.Ok

    def _clear_memos_for_analysis_restart(self) -> None:
        if self._memo_dialog is not None:
            self._memo_dialog.discard_and_close()
            self._memo_dialog = None
        if self.current_analysis_result is not None:
            self.current_analysis_result.memos.clear()
            self.current_analysis_result.handover_qa.answers = ["", "", "", "", ""]

    def _reset_for_analysis_mode_change(self) -> None:
        if self._memo_dialog is not None:
            self._memo_dialog.discard_and_close()
            self._memo_dialog = None

        self.folder_path_input.clear()
        self.folder_list_widget.clear()
        self.email_file_list_widget.clear()
        self.kakao_file_list_widget.clear()
        self.analysis_target_tabs.setCurrentIndex(self._folder_tab_index)

        # AnalysisResult owns memos, their source links, handover answers, and
        # memo.ai_result caches. Dropping it invalidates all mode-dependent data.
        self.current_analysis_result = None
        self.current_analysis_analyzed_at = None
        self._analyzed_selection_signature = None
        self.result_preview.clear()
        self._last_saved_report_fingerprint = None
        self._last_saved_word_path = None
        self._update_empty_state_labels()
        self._update_start_button_enabled()

    def _update_analysis_target_tabs_enabled(self) -> None:
        ai_mode_active = self.selected_analysis_mode == "ai"
        self.analysis_target_tabs.setTabEnabled(self._email_tab_index, ai_mode_active)
        self.analysis_target_tabs.setTabEnabled(self._kakao_tab_index, ai_mode_active)

    def _handle_analysis_target_tab_clicked(self, index: int) -> None:
        if self.selected_analysis_mode != "basic":
            return
        if index not in (self._email_tab_index, self._kakao_tab_index):
            return

        QMessageBox.warning(
            self,
            "분석 대상 선택",
            "AI 모드에서 사용 가능합니다",
        )
        self.analysis_target_tabs.setCurrentIndex(self._folder_tab_index)

    def _build_layout(self) -> None:
        central_widget = QWidget()
        central_widget.setObjectName("mainCentral")
        outer_layout = QVBoxLayout(central_widget)
        outer_layout.setContentsMargins(
            OUTER_MARGIN,
            OUTER_MARGIN,
            OUTER_MARGIN,
            OUTER_MARGIN,
        )
        outer_layout.setSpacing(0)

        app_shell_frame = QFrame()
        app_shell_frame.setObjectName("appShell")
        shadow_effect = QGraphicsDropShadowEffect(app_shell_frame)
        shadow_effect.setBlurRadius(20)
        shadow_effect.setColor(QColor(0, 0, 0, 35))
        shadow_effect.setOffset(0, 3)
        app_shell_frame.setGraphicsEffect(shadow_effect)
        outer_layout.addWidget(app_shell_frame)

        main_layout = QVBoxLayout(app_shell_frame)
        main_layout.setContentsMargins(MAIN_MARGIN, MAIN_MARGIN, MAIN_MARGIN, MAIN_MARGIN)
        main_layout.setSpacing(0)
        main_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        mode_section_layout = QVBoxLayout()
        mode_section_layout.setSpacing(MODE_SECTION_SPACING)
        mode_title = QLabel("1. 분석 모드")
        mode_title.setObjectName("sectionTitle")
        mode_title_row = QHBoxLayout()
        mode_title_row.addWidget(mode_title)
        mode_title_row.addStretch()
        mode_section_layout.addLayout(mode_title_row)
        mode_layout = QHBoxLayout()
        mode_layout.setSpacing(MODE_CARD_SPACING)
        mode_layout.addWidget(
            self._create_mode_card(
                "basic",
                "기본 모드",
                "문서, 파일 분석 → 업무 정리",
            ),
            1,
        )
        mode_layout.addWidget(
            self._create_mode_card(
                "ai",
                "AI 모드",
                "기본 모드 + GPT 분석",
            ),
            1,
        )
        mode_section_layout.addLayout(mode_layout)
        mode_actions_row = QHBoxLayout()
        mode_actions_row.setSpacing(MODE_ACTION_SPACING)
        mode_actions_row.addWidget(self.license_unlock_button, 1)
        mode_actions_row.addWidget(self.api_key_button, 1)
        mode_section_layout.addLayout(mode_actions_row)
        # Reserved for a possible future mode split (e.g. a separate mail/messenger
        # analysis mode); currently only "basic"/"ai" are selectable.
        mode_section_layout.addWidget(self.analysis_mode_notice)

        target_title = QLabel("2. 분석 대상")
        target_title.setObjectName("sectionTitle")

        self.analysis_target_tabs = _EqualWidthTabWidget()
        self.analysis_target_tabs.setObjectName("analysisTargetTabs")
        self.analysis_target_tabs.setTabBar(
            _EqualWidthTabBar(self.analysis_target_tabs)
        )

        folder_tab = QWidget()
        folder_tab_layout = QVBoxLayout(folder_tab)
        folder_row = QHBoxLayout()
        folder_row.addWidget(self.folder_path_input)
        folder_row.addWidget(self.select_folder_button)
        folder_list_actions_row = QHBoxLayout()
        folder_list_actions_row.addWidget(self.remove_selected_folders_button)
        folder_list_actions_row.addWidget(self.remove_all_folders_button)
        folder_list_actions_row.addStretch()
        folder_tab_layout.addLayout(folder_row)
        folder_tab_layout.addSpacing(10)
        folder_tab_layout.addLayout(folder_list_actions_row)
        folder_tab_layout.addSpacing(6)
        folder_tab_layout.addWidget(
            self._create_list_stack(self.folder_list_widget, self.folder_empty_label)
        )
        folder_tab_layout.addStretch()

        email_tab = QWidget()
        email_tab_layout = QVBoxLayout(email_tab)
        email_actions_row = QHBoxLayout()
        email_actions_row.addWidget(self.add_email_files_button)
        email_actions_row.addWidget(self.remove_selected_email_files_button)
        email_actions_row.addWidget(self.remove_all_email_files_button)
        email_actions_row.addStretch()
        email_tab_layout.addWidget(self.email_format_notice_label)
        email_tab_layout.addSpacing(10)
        email_tab_layout.addLayout(email_actions_row)
        email_tab_layout.addSpacing(6)
        email_tab_layout.addWidget(
            self._create_list_stack(
                self.email_file_list_widget,
                self.email_empty_label,
            )
        )
        email_tab_layout.addStretch()

        kakao_tab = QWidget()
        kakao_tab_layout = QVBoxLayout(kakao_tab)
        kakao_actions_row = QHBoxLayout()
        kakao_actions_row.addWidget(self.add_kakao_files_button)
        kakao_actions_row.addWidget(self.remove_selected_kakao_files_button)
        kakao_actions_row.addWidget(self.remove_all_kakao_files_button)
        kakao_actions_row.addStretch()
        kakao_tab_layout.addWidget(self.kakao_format_notice_label)
        kakao_tab_layout.addSpacing(10)
        kakao_tab_layout.addLayout(kakao_actions_row)
        kakao_tab_layout.addSpacing(6)
        kakao_tab_layout.addWidget(
            self._create_list_stack(
                self.kakao_file_list_widget,
                self.kakao_empty_label,
            )
        )
        kakao_tab_layout.addStretch()

        self._folder_tab_index = self.analysis_target_tabs.addTab(folder_tab, "폴더")
        self._email_tab_index = self.analysis_target_tabs.addTab(email_tab, "메일")
        self._kakao_tab_index = self.analysis_target_tabs.addTab(
            kakao_tab, "메신저(카톡 등)"
        )
        analysis_target_frame = QFrame()
        analysis_target_frame.setObjectName("analysisTargetFrame")
        analysis_target_frame_layout = QVBoxLayout(analysis_target_frame)
        analysis_target_frame_layout.setContentsMargins(0, 0, 0, 0)
        analysis_target_frame_layout.setSpacing(0)
        analysis_target_frame_layout.addWidget(self.analysis_target_tabs)

        action_separator = QFrame()
        action_separator.setObjectName("actionSeparator")
        action_separator.setFrameShape(QFrame.Shape.HLine)
        action_separator.setFixedHeight(ACTION_SEPARATOR_HEIGHT)

        action_row = QHBoxLayout()
        action_row.addWidget(self.edit_memo_button, 5)
        action_row.addSpacing(ACTION_BUTTON_SPACING)
        action_row.addWidget(self.create_rag_package_button, 2)
        action_row.addSpacing(ACTION_BUTTON_SPACING)
        action_row.addWidget(self.chatbot_button, 1)

        footer_row = QHBoxLayout()
        footer_row.setSpacing(8)
        footer_balance_spacer = QWidget()
        footer_balance_spacer.setFixedWidth(110)
        footer_row.addWidget(footer_balance_spacer)
        footer_row.addStretch()
        footer_row.addWidget(self.copyright_label)
        footer_row.addStretch()
        footer_row.addWidget(self.feedback_button)

        header_row_layout = QHBoxLayout()
        header_row_layout.addWidget(self.app_title_label)
        header_row_layout.addStretch()
        header_row_layout.addWidget(self.brand_label)
        main_layout.addLayout(header_row_layout)
        main_layout.addWidget(self.license_status_label)
        main_layout.addWidget(self.credit_balance_banner)
        main_layout.addSpacing(MAIN_SECTION_SPACING)
        main_layout.addLayout(mode_section_layout)
        main_layout.addSpacing(MAIN_SECTION_SPACING)
        main_layout.addWidget(target_title)
        main_layout.addSpacing(MODE_SECTION_SPACING)
        main_layout.addWidget(analysis_target_frame)
        main_layout.addSpacing(MODE_SECTION_SPACING)
        main_layout.addWidget(action_separator)
        main_layout.addSpacing(MODE_SECTION_SPACING)
        main_layout.addWidget(self.workflow_progress)
        main_layout.addSpacing(4)
        main_layout.addLayout(action_row)
        main_layout.addSpacing(FOOTER_SPACING)
        main_layout.addLayout(footer_row)

        self.setCentralWidget(central_widget)
        self._update_mode_card_styles()
        self._update_analysis_target_tabs_enabled()
        self.resize(WINDOW_WIDTH, max(WINDOW_MIN_HEIGHT, central_widget.sizeHint().height()))
        self._update_empty_state_labels()
        QTimer.singleShot(0, self._start_startup_credit_maintenance)
        self._credit_refresh_timer = QTimer(self)
        self._credit_refresh_timer.timeout.connect(self._refresh_credit_balance)
        self._credit_refresh_timer.start(5 * 60 * 1000)

    def _connect_signals(self) -> None:
        self.select_folder_button.clicked.connect(self._select_folder)
        self.remove_selected_folders_button.clicked.connect(
            self._remove_selected_folders
        )
        self.remove_all_folders_button.clicked.connect(self._remove_all_folders)
        self.edit_memo_button.clicked.connect(self._open_memos_for_current_analysis)
        self.create_rag_package_button.clicked.connect(self._create_rag_package)
        self.chatbot_button.clicked.connect(self._open_chatbot)
        self.feedback_button.clicked.connect(self._open_feedback_dialog)
        self.save_json_button.clicked.connect(self._save_json)
        self.license_unlock_button.clicked.connect(self._activate_license)
        self.api_key_button.clicked.connect(self._configure_api_key)
        self.add_email_files_button.clicked.connect(self._select_email_files)
        self.remove_selected_email_files_button.clicked.connect(
            self._remove_selected_email_files
        )
        self.remove_all_email_files_button.clicked.connect(
            self._remove_all_email_files
        )
        self.add_kakao_files_button.clicked.connect(self._select_kakao_files)
        self.remove_selected_kakao_files_button.clicked.connect(
            self._remove_selected_kakao_files
        )
        self.remove_all_kakao_files_button.clicked.connect(
            self._remove_all_kakao_files
        )
        self.folder_list_widget.itemClicked.connect(
            self._handle_checkable_list_item_clicked
        )
        self.email_file_list_widget.itemClicked.connect(
            self._handle_checkable_list_item_clicked
        )
        self.kakao_file_list_widget.itemClicked.connect(
            self._handle_checkable_list_item_clicked
        )
        self.analysis_target_tabs.tabBar().disabledTabClicked.connect(
            self._handle_analysis_target_tab_clicked
        )
        self._update_analysis_mode_notice()

    def _update_mode_card_styles(self, *_args: object) -> None:
        for mode_key, (card, symbol_label) in self.mode_cards.items():
            if mode_key == self.selected_analysis_mode:
                card.setProperty("selected", "true")
                symbol_label.setText("●")
                symbol_label.setProperty("selected", "true")
            else:
                card.setProperty("selected", "false")
                symbol_label.setText("○")
                symbol_label.setProperty("selected", "false")
            card.style().unpolish(card)
            card.style().polish(card)
            symbol_label.style().unpolish(symbol_label)
            symbol_label.style().polish(symbol_label)

    def _update_analysis_mode_notice(self, *_args: object) -> None:
        self.analysis_mode_notice.setText("")

    def _get_selected_analysis_mode(self) -> str:
        return self.selected_analysis_mode

    def _evaluate_startup_license_state(self) -> tuple[bool, str | None]:
        """프로그램 시작 시 인터넷 연결과 라이선스를 서버와 함께 재확인한다.

        반환값: (license_activated, lock_reason). lock_reason은 활성화 상태면
        None이고, 아니면 "not_registered" | "invalid_format" | "expired" |
        "no_internet" | "device_id_failed" | "not_found" | "server_error" |
        "other_device" | "license_terminated" 중 하나이다.
        """
        license_code = load_saved_license_code()
        if not license_code:
            return False, "not_registered"
        if not validate_license(license_code):
            return False, "invalid_format"
        if not is_license_active():
            return False, "expired"

        if not check_server_reachable():
            return False, "no_internet"

        device_id = get_device_id()
        if device_id is None:
            return False, "device_id_failed"

        status, _message = verify_license_with_server(license_code, device_id)
        if status in ("activated", "already_activated_same_device"):
            return True, None
        if status == "network_error":
            return False, "server_error"
        if status == "activated_on_other_device":
            return False, "other_device"
        if status == "license_terminated":
            return False, "license_terminated"
        if status == "not_found":
            return False, "not_found"
        return False, "server_error"

    def _activate_license(self) -> None:
        saved_license_code = load_saved_license_code()
        masked_license_code = (
            _mask_license_code(saved_license_code) if saved_license_code else ""
        )
        dialog = QInputDialog(self)
        dialog.setWindowTitle("라이선스 등록")
        dialog.setLabelText(
            "현재 등록된 라이선스가 있습니다. 변경하려면 새 라이선스 코드를\n"
            "입력하고 등록을 눌러주세요."
            if saved_license_code
            else "등록된 라이선스가 없습니다. 라이선스 코드를 입력해 주세요."
        )
        dialog.setTextValue(masked_license_code)
        dialog.setOkButtonText("등록")
        dialog.setCancelButtonText("취소")

        confirmed = dialog.exec() == QDialog.DialogCode.Accepted
        license_code = dialog.textValue()
        if not confirmed or not license_code.strip():
            return

        license_code = license_code.strip()
        if "●" in license_code or license_code == masked_license_code:
            return

        if not validate_license(license_code):
            QMessageBox.warning(
                self,
                "라이선스 등록",
                "유효하지 않거나 만료된 라이선스 코드입니다.",
            )
            return

        device_id = get_device_id()
        if device_id is None:
            QMessageBox.warning(
                self,
                "라이선스 등록",
                "PC 식별에 실패했습니다.",
            )
            return

        status, message = verify_license_with_server(license_code, device_id)
        if status in ("activated", "already_activated_same_device"):
            save_license(license_code)
            self.license_activated = is_license_active()
            self._license_lock_reason = None if self.license_activated else "expired"
            self._apply_license_lock_state(show_warning=not self.license_activated)
            if self.license_activated:
                self.credit_balance_banner.hide()
                self._refresh_credit_balance()
                QMessageBox.information(
                    self,
                    "라이선스 등록",
                    "라이선스가 등록되었습니다.",
                )
            else:
                QMessageBox.warning(
                    self,
                    "라이선스 등록",
                    "유효하지 않거나 만료된 라이선스 코드입니다.",
                )
            return

        if status == "network_error":
            QMessageBox.warning(
                self,
                "라이선스 등록",
                "서버 연결에 실패했습니다. 인터넷 연결을 확인해주세요.",
            )
            return
        if status == "activated_on_other_device":
            QMessageBox.warning(
                self,
                "라이선스 등록",
                "이미 다른 기기에서 사용 중인 라이선스입니다.",
            )
            return
        if status == "license_terminated":
            QMessageBox.warning(
                self,
                "라이선스 종료",
                f"라이선스가 종료되었습니다.\n사유: {message or '관리자에 의해 종료됨'}",
            )
            return
        if status == "not_found":
            QMessageBox.warning(
                self,
                "라이선스 등록",
                message
                or "입력하신 라이선스 코드를 찾을 수 없습니다. 코드를 다시 확인해 주세요.",
            )
            return
        QMessageBox.warning(
            self,
            "라이선스 등록",
            "알 수 없는 오류가 발생했습니다. 담당 컨설턴트에게 문의해 주세요.",
        )
        return

    def _apply_license_lock_state(self, *, show_warning: bool) -> None:
        self.license_status_label.setText(
            _LICENSE_LOCK_LABEL_TEXT.get(self._license_lock_reason, "라이선스가 만료되었습니다")
        )
        self.license_status_label.setVisible(not self.license_activated)
        if self.license_activated:
            self._set_all_buttons_enabled(True)
            self._update_start_button_enabled()
            self._update_analysis_target_tabs_enabled()
            return

        self._set_all_buttons_enabled(False)
        # 인터넷 연결/PC 식별이 안 되면 라이선스 등록 자체가 불가능하므로 등록 버튼도 잠근다.
        self.license_unlock_button.setEnabled(
            self._license_lock_reason not in ("no_internet", "device_id_failed")
        )
        if show_warning:
            QTimer.singleShot(0, self._show_license_lock_warning)

    def _show_license_lock_warning(self) -> None:
        title, message_html = _LICENSE_LOCK_DIALOG_TEXT.get(
            self._license_lock_reason, _LICENSE_LOCK_DIALOG_TEXT["expired"]
        )
        dialog = QDialog(self)
        dialog.setWindowTitle(title)
        dialog.setFixedSize(420, 170)

        message = QLabel(message_html, dialog)
        message.setAlignment(Qt.AlignmentFlag.AlignCenter)
        message.setOpenExternalLinks(True)
        message.setTextFormat(Qt.TextFormat.RichText)
        message.setWordWrap(True)

        ok_button = QPushButton("확인", dialog)
        ok_button.clicked.connect(dialog.accept)

        layout = QVBoxLayout(dialog)
        layout.addWidget(message)
        layout.addWidget(ok_button)
        dialog.exec()

    def _refresh_api_key_button_state(self) -> None:
        has_api_key = bool(self.api_key)
        self.api_key_button.setEnabled(not has_api_key)
        self.api_key_button.setToolTip(_API_KEY_ALREADY_SET_TOOLTIP if has_api_key else "")

    def _show_api_key_guide_dialog(self) -> None:
        dialog = QDialog(self)
        dialog.setWindowTitle("API 키 발급 안내")
        dialog.resize(620, 520)

        text = QTextEdit(dialog)
        text.setReadOnly(True)
        text.setPlainText(_API_KEY_GUIDE_BODY)

        copy_button = QPushButton("프롬프트 복사", dialog)
        ok_button = QPushButton("확인", dialog)
        copy_button.clicked.connect(
            lambda: QApplication.clipboard().setText(_API_KEY_GUIDE_PROMPT)
        )
        ok_button.clicked.connect(dialog.accept)

        button_row = QHBoxLayout()
        button_row.addWidget(copy_button)
        button_row.addStretch()
        button_row.addWidget(ok_button)

        layout = QVBoxLayout(dialog)
        layout.addWidget(text)
        layout.addLayout(button_row)
        dialog.exec()

    def _configure_api_key(self) -> None:
        self.api_key = load_api_key()
        if not self.api_key:
            self._show_api_key_guide_dialog()

        api_key, confirmed = QInputDialog.getText(
            self,
            "API 키 설정",
            "OpenAI API 키를 입력하세요:",
            QLineEdit.EchoMode.Password,
            _API_KEY_MASK if self.api_key else "",
        )
        if not confirmed or not api_key.strip():
            return

        api_key = api_key.strip()
        if self.api_key and api_key == _API_KEY_MASK:
            return

        save_api_key(api_key)
        self.api_key = api_key
        self._refresh_api_key_button_state()
        QMessageBox.information(
            self,
            "API 키 설정",
            "API 키가 저장되었습니다.",
        )

    def _open_chatbot(self) -> None:
        if self._chatbot_dialog is None:
            self._chatbot_dialog = ChatbotDialog(self)
        self._chatbot_dialog.show()
        self._chatbot_dialog.raise_()
        self._chatbot_dialog.activateWindow()

    def changeEvent(self, event) -> None:
        super().changeEvent(event)
        if event.type() == QEvent.Type.WindowActivate:
            QTimer.singleShot(0, self._refresh_credit_balance)

    def _refresh_credit_balance(self) -> None:
        self._start_credit_balance_worker(flush_pending=False)

    def _start_startup_credit_maintenance(self) -> None:
        self._start_credit_balance_worker(flush_pending=True)

    def _start_credit_balance_worker(self, *, flush_pending: bool) -> None:
        worker = CreditBalanceWorker(
            load_saved_license_code() or "", flush_pending=flush_pending, parent=self
        )
        worker.completed.connect(self._apply_credit_balance)
        self._track_credit_worker(worker)

    def _track_credit_worker(self, worker: QThread) -> None:
        self._credit_workers.add(worker)
        worker.finished.connect(lambda worker=worker: self._credit_workers.discard(worker))
        worker.finished.connect(worker.deleteLater)
        worker.start()

    @Slot(str, object)
    def _apply_credit_balance(self, license_code: str, balance: object) -> None:
        if license_code.strip() != (load_saved_license_code() or "").strip():
            return
        if not isinstance(balance, dict):
            self.credit_balance_banner.hide()
            return
        if not balance.get("low_balance"):
            self.credit_balance_banner.hide()
            return
        remaining = int(balance.get("balance", 0) or 0)
        granted = int(balance.get("granted_total", 0) or 0)
        percent = round((remaining / granted) * 100) if granted > 0 else 0
        self.credit_balance_banner.setText(
            f"⚠ 사용량이 얼마 남지 않았습니다 (잔여 {percent}%). 설명서 페이지에서 충전하세요."
        )
        self.credit_balance_banner.show()

    def _open_feedback_dialog(self) -> None:
        dialog = FeedbackDialog(self)
        dialog.exec()

    def _select_folder(self) -> None:
        selected_folder = QFileDialog.getExistingDirectory(
            self,
            "폴더 선택",
            self.folder_path_input.text(),
        )
        if not selected_folder:
            return

        self.folder_path_input.setText(selected_folder)
        self._add_folder_to_list(selected_folder)

    def _add_folder_to_list(self, folder_path: str) -> None:
        normalized_new_path = self._normalize_folder_path(folder_path)
        if normalized_new_path in {
            self._normalize_folder_path(existing_path)
            for existing_path in self._get_selected_folder_paths()
        }:
            return

        self._add_checkable_list_item(self.folder_list_widget, folder_path, folder_path)
        self._mark_report_unsaved()
        self._update_start_button_enabled()

    def _remove_selected_folders(self) -> None:
        if self.folder_list_widget.count() == 0 or not self._has_checked_list_item(
            self.folder_list_widget
        ):
            QMessageBox.warning(self, "폴더", "폴더를 선택하세요")
            return

        folder_paths = self._get_checked_list_item_values(self.folder_list_widget)
        if not self._confirm_delete_linked_items("폴더", folder_paths, "selected"):
            return
        self._unlink_deleted_items_from_memos("폴더", folder_paths)

        removed = False
        for index in reversed(range(self.folder_list_widget.count())):
            item = self.folder_list_widget.item(index)
            if item.data(_LIST_ITEM_CHECK_STATE_ROLE):
                self.folder_list_widget.takeItem(index)
                removed = True
        if removed:
            self._mark_report_unsaved()
            self._update_empty_state_labels()
        self._update_start_button_enabled()

    def _remove_all_folders(self) -> None:
        if self.folder_list_widget.count() == 0:
            QMessageBox.warning(self, "폴더", "폴더를 선택하세요")
            return

        folder_paths = self._get_selected_folder_paths()
        if not self._confirm_delete_linked_items("폴더", folder_paths, "all"):
            return
        self._unlink_deleted_items_from_memos("폴더", folder_paths)

        if self.folder_list_widget.count() > 0:
            self._mark_report_unsaved()
        self.folder_list_widget.clear()
        self._update_empty_state_labels()
        self._update_start_button_enabled()

    def _update_start_button_enabled(self) -> None:
        # [분석시작]이 삭제된 뒤로는 이 버튼 하나가 "필요하면 분석부터, 아니면
        # 바로 메모 팝업" 진입점이 되므로 선택 상태와 무관하게 항상 눌러야 한다.
        self.edit_memo_button.setEnabled(True)

    def _current_selection_signature(
        self,
    ) -> tuple[tuple[str, ...], tuple[str, ...], tuple[str, ...]]:
        return (
            tuple(
                sorted(
                    self._normalize_folder_path(path)
                    for path in self._get_selected_folder_paths()
                )
            ),
            tuple(
                sorted(
                    self._normalize_folder_path(path)
                    for path in self._get_selected_email_file_paths()
                )
            ),
            tuple(
                sorted(
                    self._normalize_folder_path(path)
                    for path in self._get_selected_kakao_file_paths()
                )
            ),
        )

    def _get_selected_folder_paths(self) -> list[str]:
        return [
            self.folder_list_widget.item(index).data(Qt.ItemDataRole.UserRole)
            for index in range(self.folder_list_widget.count())
        ]

    def _select_email_files(self) -> None:
        selected_files, _selected_filter = QFileDialog.getOpenFileNames(
            self,
            "메일 파일 선택",
            "",
            "메일 파일 (*.eml *.msg *.zip);;모든 파일 (*)",
        )
        for file_path in selected_files:
            self._add_email_file_to_list(file_path)

    def _add_email_file_to_list(self, file_path: str) -> None:
        normalized_new_path = self._normalize_folder_path(file_path)
        if normalized_new_path in {
            self._normalize_folder_path(existing_path)
            for existing_path in self._get_selected_email_file_paths()
        }:
            return

        self._add_checkable_list_item(self.email_file_list_widget, file_path, file_path)
        self._mark_report_unsaved()

    def _remove_selected_email_files(self) -> None:
        if self.email_file_list_widget.count() == 0 or not self._has_checked_list_item(
            self.email_file_list_widget
        ):
            QMessageBox.warning(self, "메일", "메일을 선택하세요")
            return

        email_file_paths = self._get_checked_list_item_values(self.email_file_list_widget)
        if not self._confirm_delete_linked_items("메일", email_file_paths, "selected"):
            return
        self._unlink_deleted_items_from_memos("메일", email_file_paths)

        removed = False
        for index in reversed(range(self.email_file_list_widget.count())):
            item = self.email_file_list_widget.item(index)
            if item.data(_LIST_ITEM_CHECK_STATE_ROLE):
                self.email_file_list_widget.takeItem(index)
                removed = True
        if removed:
            self._mark_report_unsaved()
            self._update_empty_state_labels()

    def _remove_all_email_files(self) -> None:
        if self.email_file_list_widget.count() == 0:
            QMessageBox.warning(self, "메일", "메일을 선택하세요")
            return

        email_file_paths = self._get_selected_email_file_paths()
        if not self._confirm_delete_linked_items("메일", email_file_paths, "all"):
            return
        self._unlink_deleted_items_from_memos("메일", email_file_paths)

        self._clear_email_files()

    def _clear_email_files(self) -> None:
        if self.email_file_list_widget.count() > 0:
            self._mark_report_unsaved()
        self.email_file_list_widget.clear()
        self._update_empty_state_labels()

    def _get_selected_email_file_paths(self) -> list[str]:
        return [
            self.email_file_list_widget.item(index).data(Qt.ItemDataRole.UserRole)
            for index in range(self.email_file_list_widget.count())
        ]

    def _select_kakao_files(self) -> None:
        selected_files, _selected_filter = QFileDialog.getOpenFileNames(
            self,
            "카카오톡 대화 파일 선택",
            "",
            "카카오톡 대화 파일 (*.txt);;모든 파일 (*)",
        )
        for file_path in selected_files:
            self._add_kakao_file_to_list(file_path)

    def _add_kakao_file_to_list(self, file_path: str) -> None:
        normalized_new_path = self._normalize_folder_path(file_path)
        if normalized_new_path in {
            self._normalize_folder_path(existing_path)
            for existing_path in self._get_selected_kakao_file_paths()
        }:
            return

        self._add_checkable_list_item(self.kakao_file_list_widget, file_path, file_path)
        self._mark_report_unsaved()

    def _remove_selected_kakao_files(self) -> None:
        if self.kakao_file_list_widget.count() == 0 or not self._has_checked_list_item(
            self.kakao_file_list_widget
        ):
            QMessageBox.warning(self, "메신저(카톡)", "메신저(카톡)을 선택하세요")
            return

        kakao_file_paths = self._get_checked_list_item_values(self.kakao_file_list_widget)
        if not self._confirm_delete_linked_items("메신저(카톡)", kakao_file_paths, "selected"):
            return
        self._unlink_deleted_items_from_memos("메신저(카톡)", kakao_file_paths)

        removed = False
        for index in reversed(range(self.kakao_file_list_widget.count())):
            item = self.kakao_file_list_widget.item(index)
            if item.data(_LIST_ITEM_CHECK_STATE_ROLE):
                self.kakao_file_list_widget.takeItem(index)
                removed = True
        if removed:
            self._mark_report_unsaved()
            self._update_empty_state_labels()

    def _remove_all_kakao_files(self) -> None:
        if self.kakao_file_list_widget.count() == 0:
            QMessageBox.warning(self, "메신저(카톡)", "메신저(카톡)을 선택하세요")
            return

        kakao_file_paths = self._get_selected_kakao_file_paths()
        if not self._confirm_delete_linked_items("메신저(카톡)", kakao_file_paths, "all"):
            return
        self._unlink_deleted_items_from_memos("메신저(카톡)", kakao_file_paths)

        self._clear_kakao_files()

    def _clear_kakao_files(self) -> None:
        if self.kakao_file_list_widget.count() > 0:
            self._mark_report_unsaved()
        self.kakao_file_list_widget.clear()
        self._update_empty_state_labels()

    def _get_selected_kakao_file_paths(self) -> list[str]:
        return [
            self.kakao_file_list_widget.item(index).data(Qt.ItemDataRole.UserRole)
            for index in range(self.kakao_file_list_widget.count())
        ]

    def _normalize_folder_path(self, folder_path: str) -> str:
        return str(Path(folder_path).resolve()).casefold()

    def _add_checkable_list_item(
        self,
        list_widget: QListWidget,
        text: str,
        data: object,
    ) -> None:
        item = QListWidgetItem(f"{_UNCHECKED_SYMBOL}{text}")
        item.setFlags(
            (item.flags() | Qt.ItemFlag.ItemIsEnabled | Qt.ItemFlag.ItemIsSelectable)
            & ~Qt.ItemFlag.ItemIsEditable
            & ~Qt.ItemFlag.ItemIsUserCheckable
        )
        item.setData(Qt.ItemDataRole.UserRole, data)
        item.setData(_LIST_ITEM_CHECK_STATE_ROLE, False)
        list_widget.addItem(item)
        self._update_empty_state_labels()

    def _handle_checkable_list_item_clicked(self, item: QListWidgetItem) -> None:
        is_checked = not bool(item.data(_LIST_ITEM_CHECK_STATE_ROLE))
        item.setData(_LIST_ITEM_CHECK_STATE_ROLE, is_checked)
        symbol = _CHECKED_SYMBOL if is_checked else _UNCHECKED_SYMBOL
        item.setText(f"{symbol}{item.data(Qt.ItemDataRole.UserRole)}")

    def _has_checked_list_item(self, list_widget: QListWidget) -> bool:
        return any(
            bool(list_widget.item(index).data(_LIST_ITEM_CHECK_STATE_ROLE))
            for index in range(list_widget.count())
        )

    def _get_checked_list_item_values(self, list_widget: QListWidget) -> list[str]:
        return [
            str(list_widget.item(index).data(Qt.ItemDataRole.UserRole))
            for index in range(list_widget.count())
            if list_widget.item(index).data(_LIST_ITEM_CHECK_STATE_ROLE)
        ]

    def _confirm_delete_linked_items(
        self,
        item_label: str,
        item_values: list[str],
        scope: str,
    ) -> bool:
        if not self._has_memo_links_to_items(item_label, item_values):
            return True

        scope_text = "선택한" if scope == "selected" else "전체"
        message_box = QMessageBox(self)
        message_box.setIcon(QMessageBox.Icon.Warning)
        message_box.setWindowTitle(f"{item_label} 삭제")
        label_text = "메일(압축파일 포함)" if item_label == "메일" else item_label
        message_box.setText(
            f"{scope_text} {label_text} 중 일부가 이미 작성된 메모에 연결되어 있습니다. "
            "삭제하면 해당 메모의 연결이 끊어집니다.\n계속하시겠습니까?"
        )
        continue_button = message_box.addButton(
            "계속 삭제",
            QMessageBox.ButtonRole.AcceptRole,
        )
        message_box.addButton("취소", QMessageBox.ButtonRole.RejectRole)
        message_box.exec()
        return message_box.clickedButton() == continue_button

    def _has_memo_links_to_items(self, item_label: str, item_values: list[str]) -> bool:
        result = self.current_analysis_result
        if result is None or not item_values:
            return False

        if item_label == "폴더":
            delete_roots = {self._normalize_folder_path(value) for value in item_values}
            for memo in result.memos:
                for linked_folder in memo.linked_folders:
                    if self._is_linked_folder_removed(linked_folder, delete_roots):
                        return True
            return False

        linked_attr = (
            "linked_emails"
            if item_label == "메일"
            else "linked_kakao_files"
        )
        deleted_items = {str(value) for value in item_values}
        if item_label == "메일":
            return any(
                self._is_email_link_removed(linked_email, deleted_items)
                for memo in result.memos
                for linked_email in memo.linked_emails
            )

        return any(
            bool(deleted_items.intersection(getattr(memo, linked_attr, [])))
            for memo in result.memos
        )

    def _unlink_deleted_items_from_memos(self, item_label: str, item_values: list[str]) -> None:
        result = self.current_analysis_result
        if result is None or not item_values:
            return

        if item_label == "폴더":
            delete_roots = {self._normalize_folder_path(value) for value in item_values}
            for memo in result.memos:
                remaining_folders = [
                    linked_folder
                    for linked_folder in memo.linked_folders
                    if not self._is_linked_folder_removed(linked_folder, delete_roots)
                ]
                if remaining_folders != memo.linked_folders:
                    memo.linked_folders = remaining_folders
            return

        deleted_items = {str(value) for value in item_values}
        if item_label == "메일":
            for memo in result.memos:
                remaining_emails = [
                    linked_email
                    for linked_email in memo.linked_emails
                    if not self._is_email_link_removed(linked_email, deleted_items)
                ]
                if remaining_emails != memo.linked_emails:
                    memo.linked_emails = remaining_emails
            return

        for memo in result.memos:
            remaining_kakao_files = [
                linked_kakao_file
                for linked_kakao_file in memo.linked_kakao_files
                if linked_kakao_file not in deleted_items
            ]
            if remaining_kakao_files != memo.linked_kakao_files:
                memo.linked_kakao_files = remaining_kakao_files

    def _is_email_link_removed(self, linked_email: str, deleted_items: set[str]) -> bool:
        if linked_email in deleted_items:
            return True
        for deleted_item in deleted_items:
            if Path(deleted_item).suffix.casefold() == ".zip" and linked_email.startswith(
                f"{deleted_item}::"
            ):
                return True
        return False

    def _is_linked_folder_removed(self, linked_folder: str, delete_roots: set[str]) -> bool:
        if not linked_folder:
            return False
        for root_path in self._get_selected_folder_paths():
            root = Path(root_path)
            relative_parts = linked_folder.split("/")
            candidates = [root.joinpath(*relative_parts)]
            if relative_parts and relative_parts[0] == root.name:
                candidates.append(root.joinpath(*relative_parts[1:]))

            for candidate in candidates:
                normalized_candidate = self._normalize_folder_path(str(candidate))
                for deleted_root in delete_roots:
                    if normalized_candidate == deleted_root or normalized_candidate.startswith(
                        f"{deleted_root}\\"
                    ):
                        return True
        return False

    def _start_analysis(self, show_complete_notice: bool = True) -> bool:
        folder_paths = self._get_selected_folder_paths()
        email_file_paths = self._get_selected_email_file_paths()
        kakao_file_paths = self._get_selected_kakao_file_paths()
        if not folder_paths and not email_file_paths and not kakao_file_paths:
            QMessageBox.warning(
                self,
                "분석 시작",
                "폴더, 메일, 메신저(카톡)를 선택하세요",
            )
            return False

        if not folder_paths:
            QMessageBox.warning(self, "분석 시작", "먼저 분석 대상 폴더를 선택하세요.")
            return False

        if any(not Path(folder_path).is_dir() for folder_path in folder_paths):
            QMessageBox.warning(
                self, "분석 시작", "선택한 폴더 중 찾을 수 없는 폴더가 있습니다."
            )
            return False

        if self.selected_analysis_mode == "ai" and not self.api_key:
            QMessageBox.warning(self, "분석 시작", "API 키를 먼저 설정해주세요.")
            return False

        if not self._confirm_clear_memos_for_analysis_restart():
            return False
        self._clear_memos_for_analysis_restart()

        # Captured now (not in _handle_analysis_succeeded) so a selection change
        # made by the user while this run is still in flight can't retroactively
        # get attributed to it — only set as _analyzed_selection_signature once
        # this exact selection has actually produced current_analysis_result.
        self._pending_analysis_selection_signature = (
            self._current_selection_signature()
        )

        self.result_preview.setPlainText("분석을 시작합니다...")
        self.edit_memo_button.setEnabled(False)
        progress_box = self._create_analysis_progress_dialog("분석 중입니다...")
        self._analysis_progress_box = progress_box
        self._show_analysis_progress_dialog(progress_box)

        worker = AnalysisWorker(
            folder_paths,
            self._get_selected_analysis_mode(),
            self,
        )
        worker.succeeded.connect(
            lambda result: self._handle_analysis_succeeded(
                result,
                show_complete_notice,
            )
        )
        worker.failed.connect(self._handle_analysis_failed)
        worker.finished.connect(worker.deleteLater)
        self._analysis_worker = worker
        self._analysis_progress_timer = QTimer(self)
        self._analysis_progress_timer.timeout.connect(QApplication.processEvents)
        self._analysis_progress_timer.start(250)
        worker.start()
        return True

    def _create_analysis_progress_dialog(self, message: str) -> QMessageBox:
        progress_box = QMessageBox(self)
        progress_box.setIcon(QMessageBox.Icon.NoIcon)
        progress_box.setWindowTitle("분석 시작")
        progress_box.setText(message)
        progress_box.setStandardButtons(QMessageBox.StandardButton.NoButton)
        progress_box.setWindowFlags(
            Qt.WindowType.Dialog | Qt.WindowType.WindowStaysOnTopHint
        )
        return progress_box

    def _show_analysis_progress_dialog(self, progress_box: QMessageBox) -> None:
        progress_box.show()
        progress_box.raise_()
        progress_box.activateWindow()
        QApplication.processEvents()

    @Slot(object)
    def _handle_analysis_succeeded(
        self,
        result: AnalysisResult,
        show_complete_notice: bool,
    ) -> None:
        existing_memos = (
            self.current_analysis_result.memos
            if self.current_analysis_result is not None
            else []
        )
        if existing_memos:
            result = replace(result, memos=existing_memos)
        self.current_analysis_result = result
        self.current_analysis_analyzed_at = datetime.now()
        self._analyzed_selection_signature = self._pending_analysis_selection_signature
        self._mark_report_unsaved()
        self.edit_memo_button.setEnabled(True)
        self.result_preview.setPlainText(self._format_analysis_result(result))
        self._auto_save_json()
        self._analysis_worker = None

        if show_complete_notice:
            if self._analysis_progress_box is not None:
                self._analysis_progress_box.setText("메모 작성 창을 여는 중...")
                QApplication.processEvents()
            self._edit_memos(before_exec=self._finish_analysis_progress)
            return

        self._finish_analysis_progress()

    @Slot(str)
    def _handle_analysis_failed(self, error_message: str) -> None:
        self._analysis_worker = None
        self._finish_analysis_progress()
        QMessageBox.critical(
            self,
            "분석 시작",
            f"분석에 실패했습니다.\n{error_message}",
        )

    def _finish_analysis_progress(self) -> None:
        if self._analysis_progress_timer is not None:
            self._analysis_progress_timer.stop()
            self._analysis_progress_timer.deleteLater()
            self._analysis_progress_timer = None
        if self._analysis_progress_box is not None:
            self._analysis_progress_box.hide()
            self._analysis_progress_box.deleteLater()
            self._analysis_progress_box = None
        if self.license_activated:
            self.edit_memo_button.setEnabled(True)

    def _open_memos_for_current_analysis(self) -> None:
        # [분석시작]이 이 버튼에 통합되었다: 이미 지금 선택 상태 그대로 분석된
        # 결과가 있으면(메모를 그대로 유지) 곧바로 메모 팝업을 열고, 그렇지
        # 않으면(최초 실행이거나 선택이 바뀌었으면) _start_analysis가 알아서
        # 백그라운드 분석을 실행한 뒤(필요 시 재분석 확인 다이얼로그를 거쳐)
        # 완료되는 대로 메모 팝업을 연다.
        if (
            self.current_analysis_result is not None
            and self._current_selection_signature()
            == self._analyzed_selection_signature
        ):
            self._edit_memos()
            return
        self._start_analysis(show_complete_notice=True)

    def _create_rag_package(self) -> None:
        api_key = load_api_key()
        if not api_key:
            QMessageBox.warning(self, "인수인계패키지 생성", "API 키를 먼저 설정해주세요.")
            return

        self._continue_create_rag_package(api_key)

    def _continue_create_rag_package(self, api_key: str) -> None:
        folder_paths = self._get_selected_folder_paths()
        email_file_paths = self._get_selected_email_file_paths()
        kakao_file_paths = self._get_selected_kakao_file_paths()
        if not folder_paths and not email_file_paths and not kakao_file_paths:
            QMessageBox.warning(
                self,
                "인수인계패키지 생성",
                "분석 대상을 먼저 선택해주세요.",
            )
            return

        scan_outcome: dict = {}
        scan_worker = CandidateScanWorker(
            folder_paths, self._get_selected_analysis_mode(), self
        )
        scan_dialog = CandidateScanProgressDialog(self)

        def handle_scan_succeeded(result: object, candidate_files: object) -> None:
            scan_outcome["status"] = "succeeded"
            scan_outcome["result"] = result
            scan_outcome["candidate_files"] = candidate_files
            scan_dialog.resolve_succeeded()

        def handle_scan_failed(message: str) -> None:
            scan_outcome["status"] = "failed"
            scan_outcome["message"] = message
            scan_dialog.resolve_failed()

        def handle_scan_canceled() -> None:
            scan_outcome.setdefault("status", "canceled")
            scan_dialog.resolve_canceled()

        scan_worker.succeeded.connect(handle_scan_succeeded)
        scan_worker.failed.connect(handle_scan_failed)
        scan_worker.canceled.connect(handle_scan_canceled)
        scan_dialog.cancel_requested.connect(scan_worker.request_cancel)
        scan_worker.finished.connect(scan_worker.deleteLater)
        self._candidate_scan_worker = scan_worker
        scan_worker.start()
        scan_dialog.exec()
        self._candidate_scan_worker = None

        if scan_outcome.get("status") != "succeeded":
            if scan_outcome.get("status") == "failed":
                QMessageBox.warning(
                    self,
                    "인수인계패키지 생성",
                    "파일을 확인하는 중 오류가 발생했습니다: "
                    f"{scan_outcome.get('message', '')}",
                )
            return

        result = scan_outcome["result"]
        if self.current_analysis_result is not None:
            result = replace(result, memos=self.current_analysis_result.memos)
        package_candidate_files = scan_outcome["candidate_files"]
        extra_size_bytes = 0
        for path in [*email_file_paths, *kakao_file_paths]:
            try:
                extra_size_bytes += Path(path).stat().st_size
            except OSError:
                pass
        extension_dialog = FileContentExtensionDialog(
            package_candidate_files,
            self,
            extra_size_bytes=extra_size_bytes,
        )
        self._active_extension_dialog = extension_dialog
        if extension_dialog.exec() != QDialog.DialogCode.Accepted:
            self._active_extension_dialog = None
            return
        self._active_extension_dialog = None
        extension_size_limits = extension_dialog.extension_size_limits()

        progress_box, progress_label = self._create_rag_package_progress_dialog(
            "예상 처리 용량 계산 중..."
        )
        self._set_all_buttons_enabled(False)
        self._rag_package_progress_box = progress_box
        self._rag_package_progress_label = progress_label
        self._show_rag_package_progress_dialog(progress_box)

        worker = CostEstimationWorker(
            result,
            folder_paths,
            email_file_paths,
            kakao_file_paths,
            extension_size_limits,
            self,
        )
        self._rag_package_context = {
            "api_key": api_key,
            "result": result,
            "folder_paths": folder_paths,
            "kakao_file_paths": kakao_file_paths,
            "extension_size_limits": extension_size_limits,
        }
        worker.succeeded.connect(self._handle_cost_estimation_succeeded)
        worker.failed.connect(self._handle_cost_estimation_failed)
        worker.canceled.connect(self._handle_cost_estimation_canceled)
        worker.progress.connect(self._handle_cost_estimation_progress)
        worker.finished.connect(worker.deleteLater)
        self._cost_estimation_worker = worker
        worker.start()

    @Slot(str, int, int)
    def _handle_cost_estimation_progress(self, phase: str, completed: int, total: int) -> None:
        if phase == "cost" and self._rag_package_progress_label is not None:
            self._rag_package_progress_label.setText(
                f"예상 처리 용량 계산 중... ({completed}/{total} 파일)"
            )

    @Slot(dict, list)
    def _handle_cost_estimation_succeeded(self, estimate: dict, parsed_emails: list) -> None:
        context = self._rag_package_context
        if context is None:
            self._finish_cost_estimation()
            return

        output_path = (
            Path.home()
            / "Desktop"
            / (
                f"{self._sanitize_package_name_part(self._get_handover_package_owner_name())}님 인수인계패키지_"
                f"{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            )
        )
        context["output_path"] = str(output_path)
        context["parsed_emails"] = parsed_emails
        size_bytes = max(0, int(estimate.get("estimated_size_bytes", 0)))
        size_gb = bytes_to_gb(size_bytes)
        context["estimated_size_bytes"] = size_bytes
        context["estimated_size_gb"] = size_gb
        if self._rag_package_progress_label is not None:
            self._rag_package_progress_label.setText("패키지 생성 요금을 확인하는 중입니다...")
        license_code = load_saved_license_code() or ""
        worker = PackageOrderWorker(license_code, size_gb, self)
        worker.completed.connect(self._handle_package_order_completed)
        worker.finished.connect(worker.deleteLater)
        self._package_order_worker = worker
        worker.start()

    @Slot(object)
    def _handle_package_order_completed(self, order_result: object) -> None:
        self._package_order_worker = None
        context = self._rag_package_context
        result = order_result if isinstance(order_result, dict) else None
        if context is None:
            return
        if result is None:
            if self._data_processing_payment_dialog is not None:
                self._data_processing_payment_dialog.show_check_failed()
            else:
                self._finish_cost_estimation()
                QMessageBox.warning(
                    self,
                    "인수인계패키지 생성",
                    "자료 처리 잔여량 서버에 연결하지 못했습니다. 잠시 후 다시 시도해주세요.",
                )
            return
        if result.get("allowed") is True:
            self._close_data_processing_payment_dialog()
            self._start_rag_package_from_context()
            return
        shortfall_gb = max(0.0, float(result.get("shortfallGb", 0.0) or 0.0))
        if self._data_processing_payment_dialog is not None:
            self._data_processing_payment_dialog.show_not_confirmed()
            return
        self._close_cost_estimation_progress()
        dialog = DataProcessingPaymentDialog(shortfall_gb, self)
        dialog.retry_requested.connect(self._retry_data_processing_check)
        dialog.cancel_requested.connect(self._cancel_data_processing_wait)
        self._data_processing_payment_dialog = dialog
        dialog.show()
        dialog.raise_()
        dialog.activateWindow()

    @Slot()
    def _retry_data_processing_check(self) -> None:
        dialog = self._data_processing_payment_dialog
        context = self._rag_package_context
        if dialog is None or context is None or self._package_order_worker is not None:
            return
        dialog.set_checking(True)
        license_code = load_saved_license_code() or ""
        requested_gb = float(context.get("estimated_size_gb", 0.0) or 0.0)
        worker = PackageOrderWorker(license_code, requested_gb, self)
        worker.completed.connect(self._handle_package_order_completed)
        worker.finished.connect(worker.deleteLater)
        self._package_order_worker = worker
        worker.start()

    @Slot()
    def _cancel_data_processing_wait(self) -> None:
        self._close_data_processing_payment_dialog()
        self._finish_cost_estimation()

    def _close_data_processing_payment_dialog(self) -> None:
        dialog = self._data_processing_payment_dialog
        self._data_processing_payment_dialog = None
        if dialog is not None:
            dialog.hide()
            dialog.deleteLater()

    def _close_cost_estimation_progress(self) -> None:
        if self._rag_package_progress_box is not None:
            self._rag_package_progress_box.hide()
            self._rag_package_progress_box.deleteLater()
            self._rag_package_progress_box = None
        self._rag_package_progress_label = None
        self._cost_estimation_worker = None

    @Slot(object)
    def _handle_package_payment_completed(self, payment_result: object) -> None:
        self._package_payment_worker = None
        result = payment_result if isinstance(payment_result, dict) else {}
        if result.get("status") == "completed":
            self._start_rag_package_from_context()
            return
        self._finish_cost_estimation()
        QMessageBox.warning(
            self,
            "인수인계패키지 생성",
            "결제가 완료되지 않아 패키지 생성이 취소되었습니다.",
        )

    def _start_rag_package_from_context(self) -> None:
        context = self._rag_package_context
        if context is None:
            return
        self._start_rag_package_worker(
            context["result"],
            context["folder_paths"],
            context["api_key"],
            context["output_path"],
            context["parsed_emails"],
            context["kakao_file_paths"],
            context["extension_size_limits"],
            reuse_progress=True,
        )

    @Slot(str)
    def _handle_cost_estimation_failed(self, error_message: str) -> None:
        self._finish_cost_estimation()
        QMessageBox.critical(
            self,
            "인수인계패키지 생성",
            f"예상 비용 계산에 실패했습니다.\n{error_message}",
        )

    @Slot()
    def _handle_cost_estimation_canceled(self) -> None:
        self._finish_cost_estimation()
        QMessageBox.information(
            self,
            "인수인계패키지 생성",
            "패키지 생성이 취소되었습니다.",
        )

    def _finish_cost_estimation(self, restore_buttons: bool = True) -> None:
        if self._rag_package_progress_box is not None:
            self._rag_package_progress_box.hide()
            self._rag_package_progress_box.deleteLater()
            self._rag_package_progress_box = None
        self._rag_package_progress_label = None
        self._cost_estimation_worker = None
        self._rag_package_context = None
        if restore_buttons:
            self._set_all_buttons_enabled(True)

    def _start_rag_package_worker(
        self,
        result: AnalysisResult,
        folder_paths: list[str],
        api_key: str,
        output_path: str,
        parsed_emails: list[dict],
        kakao_file_paths: list[str],
        extension_size_limits: dict[str, int | None],
        reuse_progress: bool = False,
    ) -> None:
        if reuse_progress and self._rag_package_progress_box is not None:
            progress_box = self._rag_package_progress_box
            progress_label = self._rag_package_progress_label
        else:
            progress_box, progress_label = self._create_rag_package_progress_dialog("진행 중...")

        worker = RagPackageWorker(
            result,
            folder_paths,
            api_key,
            output_path,
            parsed_emails,
            kakao_file_paths,
            extension_size_limits,
            self,
        )

        self._rag_package_progress_box = progress_box
        self._rag_package_progress_label = progress_label
        worker.progress.connect(self._handle_rag_package_progress)
        worker.succeeded.connect(self._handle_rag_package_succeeded)
        worker.failed.connect(self._handle_rag_package_failed)
        worker.canceled.connect(self._handle_rag_package_canceled)
        worker.finished.connect(worker.deleteLater)
        self._rag_package_worker = worker
        self._rag_package_progress_timer = QTimer(self)
        self._rag_package_progress_timer.timeout.connect(QApplication.processEvents)
        self._rag_package_progress_timer.start(250)
        self._show_rag_package_progress_dialog(progress_box)
        worker.start()

    def _create_rag_package_progress_dialog(
        self,
        message: str = "인수인계패키지 생성 중입니다...",
    ) -> tuple[QDialog, QLabel]:
        dialog = RagPackageProgressDialog(self._request_rag_package_cancel, self)
        dialog.setWindowTitle("인수인계패키지 생성")
        dialog.setWindowFlags(
            Qt.WindowType.Dialog | Qt.WindowType.WindowStaysOnTopHint
        )
        dialog.setFixedSize(620, 190)

        label = QLabel(message, dialog)
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label.setWordWrap(True)

        layout = QVBoxLayout(dialog)
        layout.addStretch()
        layout.addWidget(label)
        layout.addSpacing(18)

        banner_container = QWidget(dialog)
        banner_layout = QHBoxLayout(banner_container)
        banner_layout.setContentsMargins(0, 0, 0, 0)
        banner_layout.setSpacing(16)
        for position in ("left", "right"):
            banner_label = QLabel(banner_container)
            banner_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            banner_label.setWordWrap(True)
            banner_label.setOpenExternalLinks(False)
            banner_label.setCursor(Qt.CursorShape.PointingHandCursor)
            banner_label.setMinimumHeight(64)
            banner_label.setStyleSheet(
                "QLabel { "
                "color: #2563EB; "
                "background-color: rgba(245, 243, 255, 179); "
                "border: 1px dashed #DDD6FE; "
                "border-radius: 16px; "
                "padding: 16px 20px; "
                "}"
            )
            banner_label.linkActivated.connect(self._open_package_banner_url)
            banner_label.hide()
            banner_layout.addWidget(banner_label, 1)
            dialog.banner_labels[position] = banner_label
        banner_container.hide()
        dialog.banner_container = banner_container
        layout.addWidget(banner_container)
        layout.addStretch()

        parent_center = self.frameGeometry().center()
        dialog_rect = dialog.frameGeometry()
        dialog_rect.moveCenter(parent_center)
        dialog.move(dialog_rect.topLeft())
        return dialog, label

    def _show_rag_package_progress_dialog(self, dialog: QDialog) -> None:
        dialog.show()
        dialog.raise_()
        dialog.activateWindow()
        QApplication.processEvents()
        if isinstance(dialog, RagPackageProgressDialog):
            self._start_package_banner_worker(dialog)

    def _start_package_banner_worker(
        self, dialog: RagPackageProgressDialog
    ) -> None:
        if dialog.banner_fetch_started:
            return
        dialog.banner_fetch_started = True
        worker = PackageBannerWorker(self)
        worker.completed.connect(
            lambda banners, target=dialog: self._apply_package_banners(
                target, banners
            )
        )
        worker.finished.connect(worker.deleteLater)
        worker.finished.connect(
            lambda target=worker: self._package_banner_workers.discard(target)
        )
        self._package_banner_workers.add(worker)
        worker.start()

    @Slot(str)
    def _open_package_banner_url(self, link_url: str) -> None:
        url = QUrl(link_url)
        if url.isValid() and url.scheme().lower() in {"http", "https"}:
            QDesktopServices.openUrl(url)

    def _apply_package_banners(
        self,
        dialog: RagPackageProgressDialog,
        banners: object,
    ) -> None:
        if self._rag_package_progress_box is not dialog:
            return
        container = dialog.banner_container
        if container is None:
            return

        response = banners if isinstance(banners, dict) else {}
        visible_count = 0
        for position, label in dialog.banner_labels.items():
            banner = response.get(position)
            if not isinstance(banner, dict) or banner.get("active") is not True:
                label.hide()
                continue
            text = str(banner.get("text") or "").strip()
            link_url = str(banner.get("linkUrl") or "").strip()
            url = QUrl(link_url)
            if (
                not text
                or not url.isValid()
                or url.scheme().lower() not in {"http", "https"}
            ):
                label.hide()
                continue
            safe_text = html.escape(text)
            safe_url = html.escape(link_url, quote=True)
            label.setText(
                f'<a href="{safe_url}" style="color:#2563EB; '
                f'text-decoration:underline;"><span style="font-size:12px;">'
                f"{safe_text}</span></a>"
            )
            label.show()
            visible_count += 1
        container.setVisible(visible_count > 0)

    def _request_rag_package_cancel(self) -> None:
        if self._rag_package_progress_label is not None:
            self._rag_package_progress_label.setText("패키지 생성을 취소하는 중입니다...")
        if self._rag_package_worker is not None:
            self._rag_package_worker.request_cancel()
        if self._cost_estimation_worker is not None:
            self._cost_estimation_worker.request_cancel()
        if self._package_payment_worker is not None:
            self._package_payment_worker.requestInterruption()
            self._package_payment_worker = None
        if self._package_order_worker is not None:
            self._rag_package_context = None
        QApplication.processEvents()

    @Slot(str, int, int)
    def _handle_rag_package_progress(
        self,
        stage: str,
        current: int,
        total: int,
    ) -> None:
        if self._rag_package_progress_label is None:
            return
        if stage == "embedding":
            self._rag_package_progress_label.setText(
                f"임베딩 처리 중... ({current}/{total} 배치)"
            )
        elif stage == "resume":
            self._rag_package_progress_label.setText(
                f"이전 진행 상황에서 이어서 진행합니다 ({current + 1}번째 파일부터)"
            )
        else:
            self._rag_package_progress_label.setText(
                f"파일 처리 중... ({current}/{total} 파일)"
            )

    @Slot(str, int, int, int)
    def _handle_rag_package_succeeded(
        self,
        saved_path: str,
        failed_chunk_count: int,
        embedding_tokens: int,
        timed_out_file_count: int,
    ) -> None:
        context = self._rag_package_context or {}
        license_code = load_saved_license_code() or ""
        actual_gb = float(context.get("estimated_size_gb", 0.0) or 0.0)
        self._finish_rag_package_worker()
        if license_code and actual_gb > 0:
            worker = DataProcessingConsumeWorker(license_code, actual_gb, self)
            self._data_processing_consume_worker = worker
            worker.finished.connect(
                lambda: setattr(self, "_data_processing_consume_worker", None)
            )
            worker.finished.connect(worker.deleteLater)
            worker.start()
        message = f"패키지가 생성되었습니다: {saved_path}"
        if failed_chunk_count:
            message += f"\n\n{failed_chunk_count}개 청크는 임베딩 실패로 제외되었습니다."
        if timed_out_file_count:
            message += f"\n\n{timed_out_file_count}개 파일은 처리 시간 초과로 건너뛰었습니다."
        QMessageBox.information(
            self,
            "인수인계패키지 생성",
            message,
        )

    @Slot(str)
    def _handle_rag_package_failed(self, error_message: str) -> None:
        self._finish_rag_package_worker()
        QMessageBox.critical(
            self,
            "인수인계패키지 생성",
            f"인수인계패키지 생성에 실패했습니다.\n{error_message}",
        )

    @Slot()
    def _handle_rag_package_canceled(self) -> None:
        self._finish_rag_package_worker()
        QMessageBox.information(
            self,
            "인수인계패키지 생성",
            "패키지 생성이 취소되었습니다.",
        )

    def _finish_rag_package_worker(self) -> None:
        if self._rag_package_progress_timer is not None:
            self._rag_package_progress_timer.stop()
            self._rag_package_progress_timer.deleteLater()
            self._rag_package_progress_timer = None
        if self._rag_package_progress_box is not None:
            self._rag_package_progress_box.hide()
            self._rag_package_progress_box.deleteLater()
            self._rag_package_progress_box = None
        self._rag_package_progress_label = None
        self._set_all_buttons_enabled(True)
        self._rag_package_worker = None
        self._rag_package_context = None

    def _build_merged_analysis_result(self, folder_paths: list[str]) -> AnalysisResult:
        if len(folder_paths) == 1:
            sub_result = scan_folder(folder_paths[0])
            namespace = Path(folder_paths[0]).name or folder_paths[0]
            return self._wrap_single_folder_analysis_result(
                sub_result,
                folder_paths[0],
                namespace,
            )

        used_namespaces: set[str] = set()
        child_folder_summaries: list[ChildFolderSummary] = []
        folder_tree: list[FolderTreeNode] = []
        all_files: list[AnalyzedFile] = []
        total_folder_count = 0
        total_file_count = 0
        total_size_bytes = 0
        modified_within_7_days_count = 0
        modified_within_30_days_count = 0
        modified_within_90_days_count = 0
        error_count = 0

        for folder_path in folder_paths:
            sub_result = scan_folder(folder_path)
            namespace = self._make_unique_namespace(
                Path(folder_path).name or folder_path,
                used_namespaces,
            )

            total_folder_count += sub_result.total_folder_count
            total_file_count += sub_result.total_file_count
            total_size_bytes += sub_result.total_size_bytes
            modified_within_7_days_count += sub_result.modified_within_7_days_count
            modified_within_30_days_count += sub_result.modified_within_30_days_count
            modified_within_90_days_count += sub_result.modified_within_90_days_count
            error_count += sub_result.error_count

            child_folder_summaries.append(
                self._build_namespace_child_summary(namespace, sub_result)
            )
            folder_tree.append(
                FolderTreeNode(
                    name=namespace,
                    relative_path=namespace,
                    children=self._reprefix_folder_tree(
                        sub_result.folder_tree, namespace
                    ),
                )
            )
            all_files.extend(
                replace(file, relative_path=f"{namespace}/{file.relative_path}")
                for file in sub_result.all_files
            )

        return AnalysisResult(
            root_folder_path="; ".join(folder_paths),
            total_folder_count=total_folder_count,
            total_file_count=total_file_count,
            total_size_bytes=total_size_bytes,
            modified_within_7_days_count=modified_within_7_days_count,
            modified_within_30_days_count=modified_within_30_days_count,
            modified_within_90_days_count=modified_within_90_days_count,
            error_count=error_count,
            child_folder_summaries=child_folder_summaries,
            folder_tree=folder_tree,
            all_files=all_files,
        )

    def _wrap_single_folder_analysis_result(
        self,
        sub_result: AnalysisResult,
        folder_path: str,
        namespace: str,
    ) -> AnalysisResult:
        return AnalysisResult(
            root_folder_path=folder_path,
            total_folder_count=sub_result.total_folder_count,
            total_file_count=sub_result.total_file_count,
            total_size_bytes=sub_result.total_size_bytes,
            modified_within_7_days_count=sub_result.modified_within_7_days_count,
            modified_within_30_days_count=sub_result.modified_within_30_days_count,
            modified_within_90_days_count=sub_result.modified_within_90_days_count,
            error_count=sub_result.error_count,
            child_folder_summaries=[
                self._build_namespace_child_summary(namespace, sub_result)
            ],
            folder_tree=[
                FolderTreeNode(
                    name=namespace,
                    relative_path=namespace,
                    children=self._reprefix_folder_tree(
                        sub_result.folder_tree, namespace
                    ),
                )
            ],
            memos=sub_result.memos,
            all_files=[
                replace(file, relative_path=f"{namespace}/{file.relative_path}")
                for file in sub_result.all_files
            ],
            analysismode=sub_result.analysismode,
        )

    def _make_unique_namespace(self, base_name: str, used_namespaces: set[str]) -> str:
        candidate = base_name
        suffix = 2
        while candidate in used_namespaces:
            candidate = f"{base_name} ({suffix})"
            suffix += 1
        used_namespaces.add(candidate)
        return candidate

    def _reprefix_folder_tree(
        self,
        nodes: list[FolderTreeNode],
        namespace: str,
    ) -> list[FolderTreeNode]:
        return [
            FolderTreeNode(
                name=node.name,
                relative_path=f"{namespace}/{node.relative_path}",
                children=self._reprefix_folder_tree(node.children, namespace),
            )
            for node in nodes
        ]

    def _build_namespace_child_summary(
        self,
        namespace: str,
        sub_result: AnalysisResult,
    ) -> ChildFolderSummary:
        files = sub_result.all_files

        recent_modified_files = [
            RecentModifiedFile(
                file_name=file.file_name,
                relative_path=f"{namespace}/{file.relative_path}",
                modified_at=file.modified_at,
            )
            for file in sorted(
                files, key=lambda f: f.modified_timestamp, reverse=True
            )[:10]
        ]

        extension_counts: dict[str, int] = {}
        for file in files:
            extension = Path(file.file_name).suffix.lower() or "[no extension]"
            extension_counts[extension] = extension_counts.get(extension, 0) + 1
        extension_stats = [
            ExtensionStat(extension=extension, file_count=file_count)
            for extension, file_count in sorted(
                extension_counts.items(), key=lambda item: (-item[1], item[0])
            )[:10]
        ]

        candidate_pool = [
            file
            for file in files
            if file.size_bytes > 0
            and not file.is_hidden_or_system
            and not file.file_name.startswith("~$")
        ]
        candidate_pool.sort(
            key=lambda f: (
                Path(f.file_name).suffix.lower()
                in _DOCUMENT_EXTENSIONS_FOR_PRIORITY_HINT,
                f.modified_timestamp,
            ),
            reverse=True,
        )
        priority_review_file_candidates = [
            PriorityReviewFileCandidate(
                file_name=file.file_name,
                relative_path=f"{namespace}/{file.relative_path}",
                modified_at=file.modified_at,
                size_bytes=file.size_bytes,
            )
            for file in candidate_pool[:5]
        ]

        return ChildFolderSummary(
            folder_name=namespace,
            relative_path=namespace,
            total_folder_count=sub_result.total_folder_count,
            total_file_count=sub_result.total_file_count,
            total_size_bytes=sub_result.total_size_bytes,
            modified_within_30_days_count=sub_result.modified_within_30_days_count,
            recent_modified_files=recent_modified_files,
            extension_stats=extension_stats,
            priority_review_file_candidates=priority_review_file_candidates,
        )

    def _show_analysis_complete_notice(self) -> None:
        message_box = QMessageBox(self)
        message_box.setObjectName("analysisCompleteNotice")
        message_box.setIcon(QMessageBox.Icon.NoIcon)
        message_box.setWindowTitle("분석 완료")
        message_box.setText("분석이 완료되었습니다. 업무메모를 작성해 주세요.")
        write_memo_button = message_box.addButton(
            "업무메모작성하기",
            QMessageBox.ButtonRole.AcceptRole,
        )
        _set_button_role(write_memo_button, "primary")
        close_button = message_box.addButton(
            "닫기",
            QMessageBox.ButtonRole.RejectRole,
        )
        _set_button_role(close_button, "secondary")
        message_box.exec()

        if message_box.clickedButton() == write_memo_button:
            self._edit_memos()

    def _format_analysis_result(self, result: AnalysisResult) -> str:
        lines = [
            "분석이 완료되었습니다.",
            "",
            "[루트 폴더 전체 통계]",
            f"루트 폴더 경로: {result.root_folder_path}",
            f"전체 하위 폴더 수: {result.total_folder_count}",
            f"전체 파일 수: {result.total_file_count}",
            f"총 용량(byte): {result.total_size_bytes}",
            f"최근 7일 수정 파일 수: {result.modified_within_7_days_count}",
            f"최근 30일 수정 파일 수: {result.modified_within_30_days_count}",
            f"최근 90일 수정 파일 수: {result.modified_within_90_days_count}",
            f"접근/읽기 오류 수: {result.error_count}",
            "",
            "-" * 40,
            "",
            "[1depth 하위 폴더별 요약]",
        ]

        if not result.child_folder_summaries:
            lines.append("루트 바로 아래에 하위 폴더가 없습니다.")
        else:
            for summary in result.child_folder_summaries:
                lines.extend(self._format_child_folder_summary(summary))

        return "\n".join(lines)

    def _format_child_folder_summary(self, summary: ChildFolderSummary) -> list[str]:
        return [
            "",
            f"- 폴더명: {summary.folder_name}",
            f"  전체 하위 폴더 수: {summary.total_folder_count}",
            f"  전체 파일 수: {summary.total_file_count}",
            f"  총 용량(byte): {summary.total_size_bytes}",
            f"  최근 30일 수정 파일 수: {summary.modified_within_30_days_count}",
            "  최근 수정 파일 상위 10개:",
            *self._format_recent_modified_files(summary.recent_modified_files),
            "  확장자 통계 상위 10개:",
            *self._format_extension_stats(summary.extension_stats),
            "  가장 최근 파일 후보 상위 5개:",
            *self._format_priority_review_file_candidates(
                summary.priority_review_file_candidates
            ),
        ]

    def _format_recent_modified_files(
        self,
        recent_modified_files: list[RecentModifiedFile],
    ) -> list[str]:
        if not recent_modified_files:
            return ["    최근 수정 파일 없음"]

        return [
            f"    - {file.file_name} ({file.modified_at})"
            for file in recent_modified_files
        ]

    def _format_extension_stats(
        self,
        extension_stats: list[ExtensionStat],
    ) -> list[str]:
        if not extension_stats:
            return ["    확장자 통계 없음"]

        return [
            f"    - {stat.extension}: {stat.file_count}개"
            for stat in extension_stats
        ]

    def _format_priority_review_file_candidates(
        self,
        candidates: list[PriorityReviewFileCandidate],
    ) -> list[str]:
        if not candidates:
            return ["    가장 최근 파일 후보 없음"]

        return [
            f"    - {file.file_name} ({file.modified_at}, {file.size_bytes} bytes)"
            for file in candidates
        ]

    def _save_json(self) -> None:
        if (
            self.current_analysis_result is None
            or self.current_analysis_analyzed_at is None
        ):
            QMessageBox.warning(self, "JSON 저장", "먼저 분석을 실행하세요.")
            return
        if not self._validate_memos_for_report():
            return
        if not self._handle_recent_activity_unlinked_folders():
            return

        output_path, _selected_filter = QFileDialog.getSaveFileName(
            self,
            "JSON 저장",
            "analysis_result.json",
            "JSON Files (*.json);;All Files (*)",
        )
        if not output_path:
            return

        if Path(output_path).suffix.lower() != ".json":
            output_path = f"{output_path}.json"

        try:
            save_analysis_result_as_json(
                self.current_analysis_result,
                output_path,
                self.current_analysis_analyzed_at,
            )
        except OSError as exc:
            QMessageBox.critical(
                self,
                "JSON 저장",
                f"JSON 저장에 실패했습니다.\n{exc}",
            )
            return

        QMessageBox.information(self, "JSON 저장", "JSON 저장이 완료되었습니다.")

    def _edit_memos(
        self,
        initial_linked_folders: list[str] | None = None,
        before_exec=None,
    ) -> None:
        if self.current_analysis_result is None:
            QMessageBox.warning(self, "업무 메모 작성", "먼저 분석을 실행하세요.")
            return

        # If a MemoDialog is already open (e.g. triggered from within its own
        # save callback), bring it to front instead of opening a second one.
        if self._memo_dialog is not None and self._memo_dialog.isVisible():
            self._memo_dialog.activateWindow()
            self._memo_dialog.raise_()
            return

        email_file_paths = self._get_selected_email_file_paths()
        parsed_emails, _ = process_email_files(email_file_paths) if email_file_paths else ([], 0)

        before_fingerprint = self._build_report_state_fingerprint()
        dialog = MemoDialog(
            folder_tree=self.current_analysis_result.folder_tree,
            memos=self.current_analysis_result.memos,
            initial_linked_folders=initial_linked_folders,
            root_folder_path=self.current_analysis_result.root_folder_path,
            parsed_emails=parsed_emails,
            kakao_file_paths=self._get_selected_kakao_file_paths(),
            handover_qa=self.current_analysis_result.handover_qa,
            on_save_word=self._save_word_from_dialog,
            analysismode=self.current_analysis_result.analysismode,
            parent=self,
        )
        self._memo_dialog = dialog
        if before_exec is not None:
            dialog.show()
            dialog.raise_()
            dialog.activateWindow()
            QApplication.processEvents()
            before_exec()
        dialog.exec()
        self._memo_dialog = None
        after_fingerprint = self._build_report_state_fingerprint()
        if (
            after_fingerprint != before_fingerprint
            and self._last_saved_report_fingerprint != after_fingerprint
        ):
            self._mark_report_unsaved()
        self._auto_save_json()

    def _auto_save_json(self) -> None:
        if (
            self.current_analysis_result is None
            or self.current_analysis_analyzed_at is None
        ):
            return

        output_path = Path("output") / "analysis_result.json"
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            save_analysis_result_as_json(
                self.current_analysis_result,
                str(output_path),
                self.current_analysis_analyzed_at,
            )
        except Exception as exc:
            print(f"JSON 자동 저장 실패: {exc}")

    def _mark_report_saved(self) -> None:
        self._last_saved_report_fingerprint = self._build_report_state_fingerprint()

    def _mark_report_unsaved(self) -> None:
        self._last_saved_report_fingerprint = None
        self._last_saved_word_path = None

    def _build_report_state_fingerprint(self) -> str:
        result = self.current_analysis_result
        payload = {
            "folders": self._get_selected_folder_paths(),
            "emails": self._get_selected_email_file_paths(),
            "kakao": self._get_selected_kakao_file_paths(),
            "memos": [
                {
                    "title": memo.title,
                    "content": memo.content,
                    "linked_folders": memo.linked_folders,
                    "linked_files": memo.linked_files,
                    "linked_emails": memo.linked_emails,
                    "linked_kakao_files": memo.linked_kakao_files,
                    "ai_result_content_hash": memo.ai_result_content_hash,
                }
                for memo in (result.memos if result is not None else [])
            ],
            "handover_qa": (
                result.handover_qa.answers if result is not None else []
            ),
        }
        encoded = json.dumps(payload, ensure_ascii=False, sort_keys=True)
        return hashlib.sha256(encoded.encode("utf-8")).hexdigest()

    def _get_handover_package_owner_name(self) -> str:
        if not self._last_saved_word_path:
            return "사용자"
        try:
            from docx import Document

            document = Document(self._last_saved_word_path)
            if not document.tables:
                return "사용자"
            for row in document.tables[0].rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 4 and cells[0] == "인계자" and cells[3]:
                    return cells[3]
        except Exception:
            return "사용자"
        return "사용자"

    def _sanitize_package_name_part(self, value: str) -> str:
        sanitized = re.sub(r'[<>:"/\\|?*]+', "_", value).strip()
        return sanitized or "사용자"

    def _save_word_from_dialog(self) -> bool:
        if (
            self.current_analysis_result is None
            or self.current_analysis_analyzed_at is None
        ):
            QMessageBox.warning(self, "Word 저장", "먼저 분석을 실행하세요.")
            return False
        if not self._validate_memos_for_report():
            return False
        if not self._handle_recent_activity_unlinked_folders():
            return False
        if not self._confirm_ai_results_mode_mismatch():
            return False

        output_path, _selected_filter = QFileDialog.getSaveFileName(
            self,
            "인수인계서 저장",
            str(Path.home() / "Desktop" / "인수인계서.docx"),
            "Word Documents (*.docx);;All Files (*)",
        )
        if not output_path:
            return False

        if Path(output_path).suffix.lower() != ".docx":
            output_path = f"{output_path}.docx"

        if not self._refresh_ai_results_before_word_save(output_path, True):
            return False
        email_file_paths = self._get_selected_email_file_paths()
        parsed_emails, _ = (
            process_email_files(email_file_paths) if email_file_paths else ([], 0)
        )

        try:
            save_analysis_result_as_word(
                self.current_analysis_result,
                output_path,
                self.current_analysis_analyzed_at,
                parsed_emails,
            )
        except PermissionError:
            QMessageBox.critical(
                self,
                "인수인계서 저장",
                "저장에 실패했습니다. 같은 이름의 파일이 다른 프로그램(워드 등)에서"
                " 열려 있는지 확인 후, 닫고 다시 시도해주세요.",
            )
            return False
        except Exception as exc:
            QMessageBox.critical(
                self,
                "인수인계서 저장",
                f"인수인계서 저장에 실패했습니다.\n{exc}",
            )
            return False

        QMessageBox.information(
            self, "인수인계서 저장", "인수인계서가 저장되었습니다."
        )
        self._last_saved_word_path = output_path
        self._mark_report_saved()
        return True

    def _save_word(self) -> None:
        if (
            self.current_analysis_result is None
            or self.current_analysis_analyzed_at is None
        ):
            QMessageBox.warning(self, "Word 저장", "먼저 분석을 실행하세요.")
            return
        if not self._validate_memos_for_report():
            return
        if not self._handle_recent_activity_unlinked_folders():
            return
        if not self._confirm_ai_results_mode_mismatch():
            return

        output_path, _selected_filter = QFileDialog.getSaveFileName(
            self,
            "Word 저장",
            str(Path.home() / "Desktop" / "인수인계서.docx"),
            "Word Documents (*.docx);;All Files (*)",
        )
        if not output_path:
            return

        if Path(output_path).suffix.lower() != ".docx":
            output_path = f"{output_path}.docx"

        if not self._refresh_ai_results_before_word_save(output_path, False):
            return
        email_file_paths = self._get_selected_email_file_paths()
        parsed_emails, _ = (
            process_email_files(email_file_paths) if email_file_paths else ([], 0)
        )

        try:
            save_analysis_result_as_word(
                self.current_analysis_result,
                output_path,
                self.current_analysis_analyzed_at,
                parsed_emails,
            )
        except PermissionError:
            QMessageBox.critical(
                self,
                "Word 저장",
                "저장에 실패했습니다. 같은 이름의 파일이 다른 프로그램(워드 등)에서"
                " 열려 있는지 확인 후, 닫고 다시 시도해주세요.",
            )
            return
        except Exception as exc:
            QMessageBox.critical(
                self,
                "Word 저장",
                f"Word 저장에 실패했습니다.\n{exc}",
            )
            return

        QMessageBox.information(self, "Word 저장", "Word 저장이 완료되었습니다.")
        self._last_saved_word_path = output_path
        self._mark_report_saved()

    def _confirm_ai_results_mode_mismatch(self) -> bool:
        result = self.current_analysis_result
        if result is None or result.analysismode == "ai":
            return True
        if not any(memo.ai_result is not None for memo in result.memos):
            return True

        message_box = QMessageBox(self)
        message_box.setIcon(QMessageBox.Icon.Warning)
        message_box.setWindowTitle("인수인계서 저장")
        message_box.setText(
            "AI 분석 결과가 있는 메모가 있습니다. 현재 '기본 모드'로 저장하면"
            " AI 분석 내용(현황/주의사항 및 예상 할일)이 문서에 포함되지 않습니다."
            " 계속 진행하시겠습니까?"
        )
        save_as_ai_button = message_box.addButton(
            "AI 모드로 저장",
            QMessageBox.ButtonRole.AcceptRole,
        )
        save_as_basic_button = message_box.addButton(
            "기본 모드로 계속 저장",
            QMessageBox.ButtonRole.DestructiveRole,
        )
        cancel_button = message_box.addButton(
            "취소",
            QMessageBox.ButtonRole.RejectRole,
        )
        message_box.exec()

        clicked_button = message_box.clickedButton()
        if clicked_button == save_as_ai_button:
            self.selected_analysis_mode = "ai"
            self._update_mode_card_styles()
            self._update_analysis_target_tabs_enabled()
            self.current_analysis_result = replace(result, analysismode="ai")
            return True
        if clicked_button == save_as_basic_button:
            return True

        return clicked_button is not None and clicked_button != cancel_button

    def _refresh_ai_results_before_word_save(
        self, output_path: str, close_memo_dialog: bool
    ) -> bool | None:
        result = self.current_analysis_result
        if result is None or result.analysismode != "ai":
            return True
        if not self.license_activated or not self.api_key:
            return True

        pending_memos = [
            memo
            for memo in result.memos
            if memo.ai_result is None
            or memo.ai_result_content_hash != compute_memo_content_hash(memo)
            or not all(k in memo.ai_result for k in _REQUIRED_AI_KEYS)
        ]
        if not pending_memos:
            return True

        progress_box = QMessageBox(self)
        progress_box.setIcon(QMessageBox.Icon.NoIcon)
        progress_box.setWindowTitle("AI 분석 중")
        progress_box.setText(
            f"AI 분석 중입니다... (메모 {len(pending_memos)}개 처리 중)"
        )
        progress_box.setStandardButtons(QMessageBox.StandardButton.NoButton)

        self._set_all_buttons_enabled(False)
        progress_box.show()
        QApplication.processEvents()

        email_file_paths = self._get_selected_email_file_paths()
        parsed_emails, _ = (
            process_email_files(email_file_paths) if email_file_paths else ([], 0)
        )
        worker = ReportAiWorker(
            pending_memos,
            result.all_files,
            result.root_folder_path,
            parsed_emails,
            load_saved_license_code() or "",
            self,
        )
        worker.succeeded.connect(self._handle_report_ai_succeeded)
        worker.denied.connect(self._handle_report_ai_denied)
        worker.failed.connect(self._handle_report_ai_failed)
        worker.finished.connect(worker.deleteLater)
        self._pending_word_save = (output_path, close_memo_dialog)
        self._report_ai_progress_box = progress_box
        self._report_ai_worker = worker
        worker.start()
        return None

    @Slot(object)
    def _handle_report_ai_succeeded(self, license_code: str, balance: object) -> None:
        pending_save = self._finish_report_ai_worker()
        self._apply_credit_balance(license_code, balance)
        if pending_save is not None:
            output_path, close_memo_dialog = pending_save
            if self._save_word_to_path(output_path) and close_memo_dialog:
                if self._memo_dialog is not None:
                    self._memo_dialog.accept()

    @Slot()
    def _handle_report_ai_denied(self) -> None:
        self._finish_report_ai_worker()
        QMessageBox.warning(
            self,
            "인수인계서 저장",
            "크레딧이 부족합니다. 설명서 페이지에서 사용량을 구매해 주세요.",
        )

    @Slot(str)
    def _handle_report_ai_failed(self, error_message: str) -> None:
        self._finish_report_ai_worker()
        QMessageBox.critical(self, "인수인계서 저장", f"AI 분석에 실패했습니다.\n{error_message}")

    def _finish_report_ai_worker(self) -> tuple[str, bool] | None:
        pending_save = self._pending_word_save
        self._pending_word_save = None
        self._report_ai_worker = None
        if self._report_ai_progress_box is not None:
            self._report_ai_progress_box.hide()
            self._report_ai_progress_box.deleteLater()
            self._report_ai_progress_box = None
        self._set_all_buttons_enabled(True)
        if self._memo_dialog is not None:
            self._memo_dialog._set_all_dialog_buttons_enabled(True)
        return pending_save

    def _save_word_to_path(self, output_path: str) -> bool:
        email_file_paths = self._get_selected_email_file_paths()
        parsed_emails, _ = (
            process_email_files(email_file_paths) if email_file_paths else ([], 0)
        )
        try:
            save_analysis_result_as_word(
                self.current_analysis_result,
                output_path,
                self.current_analysis_analyzed_at,
                parsed_emails,
            )
        except Exception as exc:
            QMessageBox.critical(self, "인수인계서 저장", f"저장에 실패했습니다.\n{exc}")
            return False
        QMessageBox.information(self, "인수인계서 저장", "인수인계서가 저장되었습니다.")
        self._last_saved_word_path = output_path
        self._mark_report_saved()
        return True

    def _set_all_buttons_enabled(self, enabled: bool) -> None:
        effective_enabled = enabled and self.license_activated
        for button in (
            self.select_folder_button,
            self.remove_selected_folders_button,
            self.remove_all_folders_button,
            self.edit_memo_button,
            self.create_rag_package_button,
            self.chatbot_button,
            self.save_json_button,
            self.license_unlock_button,
            self.api_key_button,
            self.add_email_files_button,
            self.remove_selected_email_files_button,
            self.remove_all_email_files_button,
            self.add_kakao_files_button,
            self.remove_selected_kakao_files_button,
            self.remove_all_kakao_files_button,
        ):
            button.setEnabled(effective_enabled)
        if not self.license_activated:
            self.license_unlock_button.setEnabled(
                self._license_lock_reason not in ("no_internet", "device_id_failed")
            )
            return
        if enabled:
            self._refresh_api_key_button_state()

    def _validate_memos_for_report(self) -> bool:
        if self.current_analysis_result is None:
            return False

        if not self.current_analysis_result.memos:
            QMessageBox.warning(
                self,
                "보고서 저장",
                "최소 1개 이상의 업무 메모를 작성해야 보고서를 생성할 수 있습니다.",
            )
            return False

        incomplete_memos = [
            memo
            for memo in self.current_analysis_result.memos
            if not memo.title.strip() or not memo.content.strip()
        ]
        if incomplete_memos:
            QMessageBox.warning(
                self,
                "보고서 저장",
                "제목 또는 내용이 비어 있는 미완성 메모가 있습니다.",
            )
            return False

        return True

    def _handle_recent_activity_unlinked_folders(self) -> bool:
        missing_folders = self._get_recent_activity_unlinked_folders()
        if not missing_folders:
            return True

        message_box = QMessageBox(self)
        message_box.setIcon(QMessageBox.Icon.Warning)
        message_box.setWindowTitle("메모 누락 폴더 확인")
        message_box.setText(
            f"최근 활동이 있는데 메모가 없는 폴더가 {len(missing_folders)}개 있습니다."
        )
        message_box.setInformativeText(
            "\n".join(
                self._format_missing_folder_warning_line(folder)
                for folder in missing_folders
            )
        )
        open_memo_button = message_box.addButton(
            "메모 작성으로 이동",
            QMessageBox.ButtonRole.ActionRole,
        )
        save_anyway_button = message_box.addButton(
            "그대로 저장",
            QMessageBox.ButtonRole.AcceptRole,
        )
        cancel_button = message_box.addButton(
            "취소",
            QMessageBox.ButtonRole.RejectRole,
        )
        message_box.exec()

        clicked_button = message_box.clickedButton()
        if clicked_button == save_anyway_button:
            return True
        if clicked_button == open_memo_button:
            self._edit_memos(
                initial_linked_folders=[
                    folder.relative_path for folder in missing_folders
                ]
            )
            return False

        return clicked_button is not None and clicked_button != cancel_button

    def _get_recent_activity_unlinked_folders(self) -> list[ChildFolderSummary]:
        if (
            self.current_analysis_result is None
            or self.current_analysis_analyzed_at is None
        ):
            return []

        return [
            summary
            for summary in self.current_analysis_result.child_folder_summaries
            if (
                summary.modified_within_30_days_count >= 1
                and not self._is_folder_linked_to_any_memo(summary.relative_path)
            )
        ]

    def _is_folder_linked_to_any_memo(self, folder_relative_path: str) -> bool:
        if self.current_analysis_result is None:
            return False

        return any(
            linked_folder == folder_relative_path
            or linked_folder.startswith(f"{folder_relative_path}/")
            for memo in self.current_analysis_result.memos
            for linked_folder in memo.linked_folders
        )

    def _format_missing_folder_warning_line(
        self,
        folder: ChildFolderSummary,
    ) -> str:
        recent_7_days_count = self._count_recent_files_in_folder(folder, days=7)
        if recent_7_days_count > 0:
            return (
                f"{folder.folder_name} "
                f"(최근 7일 수정 파일 {recent_7_days_count}개, "
                f"최근 30일 수정 파일 {folder.modified_within_30_days_count}개)"
            )

        return (
            f"{folder.folder_name} "
            f"(최근 30일 수정 파일 {folder.modified_within_30_days_count}개)"
        )

    def _count_recent_files_in_folder(
        self,
        folder: ChildFolderSummary,
        days: int,
    ) -> int:
        if (
            self.current_analysis_result is None
            or self.current_analysis_analyzed_at is None
        ):
            return 0

        threshold = self.current_analysis_analyzed_at.timestamp() - (
            days * 24 * 60 * 60
        )
        return sum(
            1
            for file in self.current_analysis_result.all_files
            if (
                file.relative_path.startswith(f"{folder.relative_path}/")
                and file.modified_timestamp >= threshold
            )
        )

    def _show_placeholder(self, title: str, message: str) -> None:
        QMessageBox.information(self, title, message)
