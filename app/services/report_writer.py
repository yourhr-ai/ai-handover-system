import json
import re
from collections import Counter
from dataclasses import asdict
from datetime import datetime, timedelta, timezone
from difflib import SequenceMatcher
from email.utils import getaddresses, parsedate_to_datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.services.analysis_result import (
    AnalyzedFile,
    AnalysisResult,
    WorkMemo,
)
from app.services.handover_questions import HANDOVER_QUESTIONS
from app.services.email_file_handler import process_email_files
from app.services.kakao_file_handler import process_kakao_files


DOCUMENT_EXTENSIONS = {
    ".doc",
    ".docx",
    ".xls",
    ".xlsx",
    ".pdf",
    ".ppt",
    ".pptx",
    ".hwp",
}
SUPPORTING_EXTENSIONS = {".tmp", ".log"}
SMALL_FILE_SIZE_BYTES = 1024
PRIORITY_REVIEW_FILE_LIMIT = 5
PRIORITY_REVIEW_DEDUPLICATION_CANDIDATE_LIMIT = 30
EXTENSION_GROUPS = [
    (
        "사무문서",
        {
            "doc",
            "docx",
            "hwp",
            "hwpx",
            "xls",
            "xlsx",
            "xlsm",
            "ppt",
            "pptx",
            "pdf",
        },
    ),
    (
        "이미지/텍스트",
        {
            "jpg",
            "jpeg",
            "png",
            "gif",
            "bmp",
            "webp",
            "svg",
            "tiff",
            "txt",
            "md",
            "csv",
            "json",
            "xml",
            "yaml",
            "yml",
        },
    ),
    ("디자인", {"psd", "ai", "xd", "fig", "sketch", "indd", "eps"}),
    (
        "코드/개발",
        {
            "py",
            "js",
            "ts",
            "jsx",
            "tsx",
            "java",
            "c",
            "cpp",
            "cs",
            "go",
            "rs",
            "html",
            "css",
            "sql",
            "sh",
        },
    ),
    ("CAD/3D", {"dwg", "dxf", "skp", "3ds", "obj", "stl"}),
    (
        "영상/오디오",
        {"mp4", "mov", "avi", "mkv", "wmv", "mp3", "wav", "aac", "flac"},
    ),
    ("기타", set()),
]


def save_analysis_result_as_json(
    result: AnalysisResult,
    output_path: str,
    analyzed_at: datetime,
) -> None:
    report = _build_report(result, analyzed_at)
    Path(output_path).write_text(
        json.dumps(report, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def save_analysis_result_as_word(
    result: AnalysisResult,
    output_path: str,
    analyzed_at: datetime,
    parsed_emails: list[dict] | None = None,
) -> None:
    document = Document()
    _set_default_font(document)

    title = document.add_heading("인수인계서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_signature_table(document)

    _add_handover_qa_section(document, result, parsed_emails)

    if not result.memos:
        document.add_paragraph("작성된 업무 메모 없음")
    else:
        for memo in result.memos:
            _add_memo_block(
                document,
                memo,
                result.all_files,
                result.analysismode,
                parsed_emails,
            )

    document.add_page_break()
    _add_manual_heading(document, "문서 통계", font_size=16, space_before=18)
    _add_document_statistics_table(document, result.all_files, analyzed_at)

    _add_end_of_document_page(document)

    document.save(output_path)


_WORK_TERMS = {
    "업무", "회의", "계약", "견적", "납기", "일정", "보고", "승인", "고객",
    "자료", "파일", "프로젝트", "요청", "확인", "수정", "발송", "매출", "급여",
}
_SELF_NAMES = {"나", "내", "본인", "me", "myself", "사용자"}


def collect_key_contacts(
    parsed_emails: list[dict] | None,
    kakao_messages: list[dict] | None,
) -> list[dict]:
    """Return the ten most relevant contacts across email and messenger data.

    발신자/수신자 전체를 있는 그대로 집계하며, 추정된 본인 주소도 순위와
    상위 10명 제한에 동일하게 포함한다. 본인 추정 결과는 표시에만 사용한다.
    Messenger exports still exclude explicit self labels (e.g. "나", "본인")
    since that identification is unambiguous. Business keywords affect ranking,
    but not the period counts. Group chats are never filtered out; every
    individual sender in a group contributes separately.
    """
    now = datetime.now(timezone.utc)
    email_rows: list[tuple[list[tuple[str, str]], list[tuple[str, str]], datetime | None, str]] = []

    for email in parsed_emails or []:
        senders = _parse_email_contacts(str(email.get("sender") or ""))
        recipients = _parse_email_contacts(str(email.get("recipient") or ""))
        email_rows.append(
            (
                senders,
                recipients,
                _parse_contact_datetime(email.get("date")),
                f"{email.get('subject') or ''} {email.get('body') or ''}",
            )
        )

    self_contact_guess = _guess_self_contact(email_rows)
    stats: dict[str, dict] = {}

    for senders, recipients, occurred_at, text in email_rows:
        for name, address in [*senders, *recipients]:
            _record_contact(stats, name or address, address, occurred_at, text, now)

    for message in kakao_messages or []:
        name = str(message.get("sender") or "").strip()
        if not name or _is_self_contact(name, name):
            continue
        occurred_at = _parse_contact_datetime(
            f"{message.get('date') or ''} {message.get('time') or ''}"
        )
        _record_contact(
            stats,
            name,
            "메신저",
            occurred_at,
            str(message.get("message") or ""),
            now,
        )

    ranked = sorted(
        stats.values(),
        key=lambda item: (
            item.pop("_work_score", 0),
            item["freq_1m"],
            item["freq_3m"],
            item["freq_6m"],
        ),
        reverse=True,
    )
    for item in ranked:
        item["is_self_guess"] = (
            self_contact_guess is not None
            and item["contact"].casefold() == self_contact_guess
        )
    return ranked[:10]


def _parse_email_contacts(value: str) -> list[tuple[str, str]]:
    contacts = []
    for name, address in getaddresses([value]):
        address = address.strip()
        if address:
            contacts.append((name.strip(), address))
    return contacts


def _guess_self_contact(
    email_rows: list[
        tuple[list[tuple[str, str]], list[tuple[str, str]], datetime | None, str]
    ],
) -> str | None:
    sender_counts: Counter[str] = Counter()
    recipient_counts: Counter[str] = Counter()
    for senders, recipients, _, _ in email_rows:
        sender_counts.update(address.casefold() for _, address in senders)
        recipient_counts.update(address.casefold() for _, address in recipients)

    candidates = sender_counts.keys() & recipient_counts.keys()
    if not candidates:
        return None

    scores = {
        address: (
            min(sender_counts[address], recipient_counts[address]),
            sender_counts[address] + recipient_counts[address],
        )
        for address in candidates
    }
    best_score = max(scores.values())
    best_candidates = [
        address for address, score in scores.items() if score == best_score
    ]
    return best_candidates[0] if len(best_candidates) == 1 else None


def _is_self_contact(name: str, identifier: str) -> bool:
    return name.strip().casefold() in _SELF_NAMES or identifier.strip().casefold() in _SELF_NAMES


def _parse_contact_datetime(value: object) -> datetime | None:
    text = str(value or "").strip()
    if not text:
        return None
    try:
        parsed = parsedate_to_datetime(text)
    except (TypeError, ValueError, OverflowError):
        parsed = None
    if parsed is None:
        normalized = re.sub(r"\s*(오전|오후)\s*(\d{1,2}):(\d{2})", _replace_korean_time, text)
        try:
            parsed = datetime.fromisoformat(normalized.strip())
        except ValueError:
            try:
                parsed = datetime.strptime(normalized.strip(), "%Y-%m-%d %H:%M")
            except ValueError:
                return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=timezone.utc)
    return parsed.astimezone(timezone.utc)


def _replace_korean_time(match: re.Match) -> str:
    hour = int(match.group(2)) % 12
    if match.group(1) == "오후":
        hour += 12
    return f" {hour:02d}:{match.group(3)}"


def _record_contact(
    stats: dict[str, dict],
    name: str,
    contact: str,
    occurred_at: datetime | None,
    text: str,
    now: datetime,
) -> None:
    key = contact.casefold() if contact != "메신저" else f"messenger:{name.casefold()}"
    item = stats.setdefault(
        key,
        {
            "name": name,
            "contact": contact,
            "freq_1m": 0,
            "freq_3m": 0,
            "freq_6m": 0,
            "_work_score": 0,
        },
    )
    if occurred_at is not None:
        age = now - occurred_at
        if timedelta(0) <= age <= timedelta(days=183):
            item["freq_6m"] += 1
            if age <= timedelta(days=92):
                item["freq_3m"] += 1
            if age <= timedelta(days=31):
                item["freq_1m"] += 1
    lowered = text.casefold()
    if any(term in lowered for term in _WORK_TERMS):
        item["_work_score"] += 1


def _add_handover_qa_section(
    document: Document,
    result: AnalysisResult,
    parsed_emails: list[dict] | None,
) -> None:
    answers = (result.handover_qa.answers + ["", "", "", "", ""])[:5]
    answered_questions = [
        (question, answer)
        for question, answer in zip(HANDOVER_QUESTIONS, answers)
        if answer.strip()
    ]
    if not answered_questions:
        return

    document.add_paragraph()
    _add_manual_heading(document, "주요사항", font_size=16, space_before=18)
    for question, answer in answered_questions:
        _add_manual_heading(
            document,
            str(question["title"]),
            font_size=13,
            space_before=12,
        )
        _add_multiline_paragraph(document, answer)

    if result.analysismode == "ai":
        kakao_paths = list(
            dict.fromkeys(
                path
                for memo in result.memos
                for path in memo.linked_kakao_files
            )
        )
        kakao_messages, _ = process_kakao_files(kakao_paths) if kakao_paths else ([], 0)
        contacts = collect_key_contacts(parsed_emails, kakao_messages)
        _add_key_contacts_table(document, contacts)

    _add_separator(document)

def _add_multiline_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    lines = text.split("\n")
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        paragraph.add_run(line)


def _add_key_contacts_table(document: Document, contacts: list[dict]) -> None:
    _add_manual_heading(document, "주요 관계자 및 연락처", font_size=13, space_before=12)
    table = document.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    name_header = table.cell(0, 0).merge(table.cell(1, 0))
    contact_header = table.cell(0, 1).merge(table.cell(1, 1))
    frequency_header = table.cell(0, 2).merge(table.cell(0, 4))
    name_header.text = "이름"
    contact_header.text = "연락 방법"
    frequency_header.text = "연락 빈도"
    table.cell(1, 2).text = "최근 1개월"
    table.cell(1, 3).text = "최근 3개월"
    table.cell(1, 4).text = "최근 6개월"
    for contact in contacts:
        row = table.add_row()
        _set_row_cells(
            row,
            [
                contact["name"],
                contact["contact"],
                contact["freq_1m"],
                contact["freq_3m"],
                contact["freq_6m"],
            ],
        )
        if contact.get("is_self_guess", False):
            label_run = row.cells[0].paragraphs[0].add_run()
            label_run.add_break()
            label_run.add_text("(본인 추정)")
            label_run.font.size = Pt(9)
    if not contacts:
        row = table.add_row()
        row.cells[0].merge(row.cells[4]).text = "집계할 연락 기록이 없습니다."
    _center_table_cells(table)


def _add_memo_block(
    document: Document,
    memo: WorkMemo,
    all_files: list[AnalyzedFile],
    analysismode: str,
    parsed_emails: list[dict] | None = None,
) -> None:
    document.add_paragraph()
    _add_manual_heading(document, memo.title, font_size=16, space_before=18)
    document.add_paragraph(memo.content)

    if analysismode == "ai":
        _add_ai_result_sections(document, memo)

    group_paths = _get_group_paths_for_memo(memo)
    if not group_paths:
        document.add_paragraph("없음")
    else:
        for group_path in sorted(group_paths):
            _add_manual_heading(
                document,
                f"{_get_folder_name(group_path)} 폴더",
                font_size=13,
                space_before=12,
            )
            document.add_paragraph("가장 최근 파일")
            priority_candidates = _get_recent_files_for_top_level_folder(
                all_files,
                group_path,
                memo.linked_files,
                linked_folders=memo.linked_folders,
            )
            _add_file_location_table(document, priority_candidates)

    _add_recent_emails_section(document, memo, parsed_emails)
    _add_recent_kakao_section(document, memo)
    _add_separator(document)


def _add_ai_result_sections(document: Document, memo: WorkMemo) -> None:
    # The AI result is now computed/cached by the caller (main_window.py, at
    # Word-save time) - this only renders whatever ended up on the memo.
    ai_result = memo.ai_result
    print(
        f"[DBG4 render] title={memo.title!r}"
        f" ai_result_is_none={ai_result is None}"
        f" email_summary={str(ai_result.get('email_summary', '') if ai_result else '')!r:.80}"
        f" kakao_summary={str(ai_result.get('kakao_summary', '') if ai_result else '')!r:.80}"
    )
    if ai_result is None:
        document.add_paragraph("AI 분석 실패 - 기본 모드로 표시됩니다")
        return

    _add_manual_heading(document, "현황", font_size=13, space_before=12)
    _add_paragraph_with_source_text(document, ai_result.get("status_summary", ""))

    email_summary = str(ai_result.get("email_summary", "")).strip()
    if memo.linked_emails:
        _add_manual_heading(
            document, "메일에서 확인된 내용", font_size=13, space_before=12
        )
        _add_paragraph_with_source_text(
            document,
            email_summary
            or "선택된 메일에서 이 업무와 직접 관련된 내용은 확인되지 않습니다."
        )

    kakao_summary = str(ai_result.get("kakao_summary", "")).strip()
    if memo.linked_kakao_files:
        _add_manual_heading(
            document, "메신저(카톡)에서 확인된 사항", font_size=13, space_before=12
        )
        _add_paragraph_with_source_text(
            document,
            kakao_summary
            or "선택된 대화에서 이 업무와 직접 관련된 내용은 확인되지 않습니다."
        )

    # Placed after the per-source sections (메일/카톡) and before 주의사항: it
    # reads naturally as "here's each source individually, here's how they
    # connect, here's what to do about it" - the precautions/next-steps
    # section then follows on from the synthesis instead of coming before it.
    _add_manual_heading(
        document, "자료 간 연결 및 타임라인", font_size=13, space_before=12
    )
    _add_markdown_bullet_list(document, ai_result.get("cross_check", ""))

    _add_manual_heading(document, "주의사항 및 예상 할일", font_size=13, space_before=12)
    _add_markdown_bullet_list(document, ai_result.get("precautions_and_tasks", ""))


def get_combined_priority_review_files_for_memo(
    all_files: list[AnalyzedFile],
    memo: WorkMemo,
) -> list[AnalyzedFile]:
    combined_files: list[AnalyzedFile] = []
    seen_relative_paths: set[str] = set()
    for group_path in sorted(_get_group_paths_for_memo(memo)):
        for file in _get_priority_review_files_for_top_level_folder(
            all_files,
            group_path,
            memo.linked_files,
            linked_folders=memo.linked_folders,
        ):
            if file.relative_path not in seen_relative_paths:
                seen_relative_paths.add(file.relative_path)
                combined_files.append(file)

    return combined_files


def get_recent_files_for_memo(
    all_files: list[AnalyzedFile],
    memo: WorkMemo,
) -> list[AnalyzedFile]:
    combined_files: list[AnalyzedFile] = []
    seen_relative_paths: set[str] = set()
    for group_path in sorted(_get_group_paths_for_memo(memo)):
        for file in _get_recent_files_for_top_level_folder(
            all_files,
            group_path,
            memo.linked_files,
            linked_folders=memo.linked_folders,
        ):
            if file.relative_path not in seen_relative_paths:
                seen_relative_paths.add(file.relative_path)
                combined_files.append(file)

    return combined_files


_RECENT_EMAIL_LIMIT = 5
_RECENT_KAKAO_LIMIT = 5


def _add_recent_emails_section(
    document: Document,
    memo: WorkMemo,
    parsed_emails: list[dict] | None = None,
) -> None:
    if not memo.linked_emails:
        return
    emails = _get_linked_parsed_emails(memo, parsed_emails)
    emails_sorted = sorted(emails, key=lambda e: e.get("date") or "", reverse=True)
    top = emails_sorted[:_RECENT_EMAIL_LIMIT]
    _add_manual_heading(document, "가장 최근 메일 5개", font_size=13, space_before=12)
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    table.style = "Table Grid"
    _set_row_cells(table.rows[0], ["발신자", "메일 제목", "발송일시"])
    for email in top:
        sender = email.get("sender") or ""
        subject = email.get("subject") or "(제목 없음)"
        date = email.get("date") or ""
        _set_row_cells(table.add_row(), [sender, subject, date])
    _center_table_cells(table)


def _get_linked_parsed_emails(
    memo: WorkMemo,
    parsed_emails: list[dict] | None,
) -> list[dict]:
    if parsed_emails is None:
        emails, _ = process_email_files(memo.linked_emails)
        return emails

    linked_set = set(memo.linked_emails)
    return [
        email
        for email in parsed_emails
        if email.get("source_file", "") in linked_set
    ]


def _kakao_datetime_sort_key(msg: dict) -> tuple[str, int]:
    date = msg.get("date") or ""
    parts = (msg.get("time") or "").split()
    try:
        ampm = parts[0] if parts else ""
        h, m = (parts[1] if len(parts) > 1 else "0:00").split(":")
        hour, minute = int(h), int(m)
        if ampm == "오후" and hour != 12:
            hour += 12
        elif ampm == "오전" and hour == 12:
            hour = 0
    except (ValueError, IndexError):
        hour, minute = 0, 0
    return (date, hour * 60 + minute)


def _add_recent_kakao_section(document: Document, memo: WorkMemo) -> None:
    if not memo.linked_kakao_files:
        return
    messages, _ = process_kakao_files(memo.linked_kakao_files)
    if not messages:
        return
    messages_sorted = sorted(messages, key=_kakao_datetime_sort_key, reverse=True)
    top = messages_sorted[:_RECENT_KAKAO_LIMIT]
    _add_manual_heading(
        document, "가장 최근 대화 목록 5개", font_size=13, space_before=12
    )
    for msg in top:
        date = msg.get("date") or ""
        time = msg.get("time") or ""
        sender = msg.get("sender") or ""
        message = msg.get("message") or ""
        document.add_paragraph(
            f"{date} {time} - {sender}: {message}", style="List Bullet"
        )

def _add_markdown_bullet_list(document: Document, bullet_items: str | list) -> None:
    raw_items = bullet_items if isinstance(bullet_items, list) else bullet_items.splitlines()
    items = _join_multiline_source_items(raw_items)
    for item in items:
        stripped_item = str(item).strip()
        if not stripped_item:
            continue
        if stripped_item.startswith("- ") or stripped_item.startswith("* "):
            stripped_item = stripped_item[2:].strip()
        if stripped_item:
            _add_paragraph_with_source_text(
                document, stripped_item, style="List Bullet"
            )


_SOURCE_TEXT_PATTERN = re.compile(r"\(\s*출처\s*:\s*")


def _find_source_text_spans(source_text: str) -> tuple[list[tuple[int, int]], bool]:
    spans: list[tuple[int, int]] = []
    search_position = 0
    while match := _SOURCE_TEXT_PATTERN.search(source_text, search_position):
        depth = 1
        position = match.end()
        while position < len(source_text) and depth:
            if source_text[position] == "(":
                depth += 1
            elif source_text[position] == ")":
                depth -= 1
            position += 1
        if depth:
            return spans, True
        spans.append((match.start(), position))
        search_position = position
    return spans, False


def _join_multiline_source_items(items: list | tuple) -> list[str]:
    joined_items: list[str] = []
    pending_lines: list[str] = []
    for item in items:
        line = str(item).strip()
        if not line:
            continue
        pending_lines.append(line)
        combined = " ".join(pending_lines)
        _, has_unclosed_source = _find_source_text_spans(combined)
        if not has_unclosed_source:
            joined_items.append(combined)
            pending_lines.clear()
    if pending_lines:
        joined_items.append(" ".join(pending_lines))
    return joined_items


def _add_paragraph_with_source_text(
    document: Document,
    text: object,
    style: str | None = None,
) -> None:
    paragraph = document.add_paragraph(style=style)
    source_text = str(text or "")
    cursor = 0
    source_spans, _ = _find_source_text_spans(source_text)
    for start, end in source_spans:
        if start > cursor:
            run = paragraph.add_run(source_text[cursor:start])
            run.font.size = Pt(11)
        source_run = paragraph.add_run(source_text[start:end])
        source_run.font.size = Pt(9)
        cursor = end
    if cursor < len(source_text):
        run = paragraph.add_run(source_text[cursor:])
        run.font.size = Pt(11)



def _build_report(result: AnalysisResult, analyzed_at: datetime) -> dict[str, Any]:
    linked_top_level_folders = _get_linked_top_level_folders(result)

    return {
        "analyzed_at": analyzed_at.isoformat(timespec="seconds"),
        "analysismode": result.analysismode,
        "root_folder_path": result.root_folder_path,
        "root_summary": {
            "total_folder_count": result.total_folder_count,
            "total_file_count": result.total_file_count,
            "total_size_bytes": result.total_size_bytes,
            "modified_within_7_days_count": result.modified_within_7_days_count,
            "modified_within_30_days_count": result.modified_within_30_days_count,
            "modified_within_90_days_count": result.modified_within_90_days_count,
        },
        "error_count": result.error_count,
        "handover_qa": asdict(result.handover_qa),
        "memos": [_build_memo_report(memo, result, analyzed_at) for memo in result.memos],
        "sharedfolders": _build_shared_folders(result),
        "unlinkedfolders": [
            _build_unlinked_folder_report(summary, result, analyzed_at)
            for summary in result.child_folder_summaries
            if summary.relative_path not in linked_top_level_folders
        ],
        "child_folders": [
            {
                "folder_name": summary.folder_name,
                "relative_path": summary.relative_path,
                "summary": {
                    "total_folder_count": summary.total_folder_count,
                    "total_file_count": summary.total_file_count,
                    "total_size_bytes": summary.total_size_bytes,
                    "modified_within_30_days_count": (
                        summary.modified_within_30_days_count
                    ),
                },
                "recent_modified_files": [
                    asdict(file) for file in summary.recent_modified_files
                ],
                "extension_stats": [
                    asdict(stat) for stat in summary.extension_stats
                ],
                "priority_review_file_candidates": [
                    asdict(file)
                    for file in summary.priority_review_file_candidates
                ],
            }
            for summary in result.child_folder_summaries
        ],
    }


def _build_memo_report(
    memo: WorkMemo,
    result: AnalysisResult,
    analyzed_at: datetime,
) -> dict[str, Any]:
    return {
        "title": memo.title,
        "content": memo.content,
        "createdat": memo.createdat,
        "updatedat": memo.updatedat,
        "linked_folders": memo.linked_folders,
        "linked_files": memo.linked_files,
        "linked_emails": memo.linked_emails,
        "linked_kakao_files": memo.linked_kakao_files,
        "linked_folder_stats": [
            _build_linked_folder_stats(relative_path, result, analyzed_at)
            for relative_path in memo.linked_folders
        ],
        "priorityfilesbytopfolder": _build_priority_files_by_top_folder(
            memo,
            result,
        ),
    }


def _build_priority_files_by_top_folder(
    memo: WorkMemo,
    result: AnalysisResult,
) -> dict[str, list[dict[str, str]]]:
    priority_files_by_top_folder: dict[str, list[dict[str, str]]] = {}
    files_by_relative_path = _get_all_files_by_relative_path(result)

    for group_path in sorted(_get_group_paths_for_memo(memo)):
        priority_files_by_top_folder[group_path] = [
            {
                "file_name": file.file_name,
                "relative_path": file.relative_path,
                "modified_at": file.modified_at,
            }
            for file in _get_recent_files_for_top_level_folder(
                result.all_files,
                group_path,
                memo.linked_files,
                files_by_relative_path,
                linked_folders=memo.linked_folders,
            )
        ]

    return priority_files_by_top_folder


def _get_all_files_by_relative_path(
    result: AnalysisResult,
) -> dict[str, AnalyzedFile]:
    # Cached directly on the (frozen) result instance so repeated report
    # builds for the same analysis don't rebuild this for every memo/folder.
    cache = getattr(result, "_all_files_by_relative_path_cache", None)
    if cache is None:
        cache = {file.relative_path: file for file in result.all_files}
        object.__setattr__(result, "_all_files_by_relative_path_cache", cache)
    return cache


def _build_linked_folder_stats(
    relative_path: str,
    result: AnalysisResult,
    analyzed_at: datetime,
) -> dict[str, Any]:
    summary = _find_top_level_summary(relative_path, result)
    if summary is None:
        return {
            "relative_path": relative_path,
            "source_1depth_relative_path": None,
            "file_count": 0,
            "total_size_bytes": 0,
            "extension_counts": {},
            "modified_within_7_days_count": 0,
            "modified_within_30_days_count": 0,
            "modified_within_90_days_count": 0,
        }

    return {
        "relative_path": relative_path,
        "source_1depth_relative_path": summary.relative_path,
        "file_count": summary.total_file_count,
        "total_size_bytes": summary.total_size_bytes,
        "extension_counts": {
            stat.extension: stat.file_count for stat in summary.extension_stats
        },
        "modified_within_7_days_count": _count_files_modified_within_days(
            result,
            summary.relative_path,
            analyzed_at,
            7,
        ),
        "modified_within_30_days_count": summary.modified_within_30_days_count,
        "modified_within_90_days_count": _count_files_modified_within_days(
            result,
            summary.relative_path,
            analyzed_at,
            90,
        ),
    }


def _build_shared_folders(result: AnalysisResult) -> dict[str, list[str]]:
    folder_to_titles: dict[str, list[str]] = {}
    for memo in result.memos:
        memo_title = memo.title
        for relative_path in memo.linked_folders:
            folder_to_titles.setdefault(relative_path, []).append(memo_title)

    return {
        relative_path: titles
        for relative_path, titles in folder_to_titles.items()
        if len(titles) >= 2
    }


def _build_unlinked_folder_report(
    summary: Any,
    result: AnalysisResult,
    analyzed_at: datetime,
) -> dict[str, Any]:
    return {
        "folder_name": summary.folder_name,
        "relative_path": summary.relative_path,
        "file_count": summary.total_file_count,
        "modified_within_30_days_count": summary.modified_within_30_days_count,
        "modified_within_90_days_count": _count_files_modified_within_days(
            result,
            summary.relative_path,
            analyzed_at,
            90,
        ),
    }


def _get_linked_top_level_folders(result: AnalysisResult) -> set[str]:
    linked_top_level_folders: set[str] = set()
    top_level_paths = {
        summary.relative_path for summary in result.child_folder_summaries
    }
    for memo in result.memos:
        for relative_path in memo.linked_folders:
            top_level_path = relative_path.split("/", 1)[0]
            if top_level_path in top_level_paths:
                linked_top_level_folders.add(top_level_path)

    return linked_top_level_folders


def _find_top_level_summary(relative_path: str, result: AnalysisResult) -> Any:
    top_level_path = relative_path.split("/", 1)[0]
    for summary in result.child_folder_summaries:
        if summary.relative_path == top_level_path:
            return summary

    return None


def _count_files_modified_within_days(
    result: AnalysisResult,
    folder_relative_path: str,
    analyzed_at: datetime,
    days: int,
) -> int:
    threshold = analyzed_at.timestamp() - (days * 24 * 60 * 60)
    return sum(
        1
        for file in result.all_files
        if (
            file.relative_path.startswith(f"{folder_relative_path}/")
            and file.modified_timestamp >= threshold
        )
    )


def _add_document_statistics_table(
    document: Document,
    all_files: list[AnalyzedFile],
    analyzed_at: datetime,
) -> None:
    group_stats = _build_extension_group_stats(all_files, analyzed_at)
    active_groups = [
        group_name
        for group_name, _extensions in EXTENSION_GROUPS
        if group_stats[group_name]["total"] > 0
    ]
    if not active_groups:
        document.add_paragraph("없음")
        return

    table = document.add_table(rows=4, cols=len(active_groups) + 1)
    table.style = "Table Grid"
    _set_row_cells(table.rows[0], ["구분", *active_groups])
    _set_row_cells(
        table.rows[1],
        [
            "전체 파일 수",
            *[_format_count(group_stats[group]["total"]) for group in active_groups],
        ],
    )
    _set_row_cells(
        table.rows[2],
        [
            "10일내 작업",
            *[_format_count(group_stats[group]["within_10"]) for group in active_groups],
        ],
    )
    _set_row_cells(
        table.rows[3],
        [
            "11~30일내 작업",
            *[
                _format_count(group_stats[group]["over_10_within_30"])
                for group in active_groups
            ],
        ],
    )
    _center_table_cells(table)


def _build_extension_group_stats(
    all_files: list[AnalyzedFile],
    analyzed_at: datetime,
) -> dict[str, dict[str, int]]:
    group_stats = {
        group_name: {
            "total": 0,
            "within_10": 0,
            "over_10_within_30": 0,
            "over_30": 0,
        }
        for group_name, _extensions in EXTENSION_GROUPS
    }
    ten_days = 10 * 24 * 60 * 60
    thirty_days = 30 * 24 * 60 * 60

    for file in all_files:
        group_name = _get_extension_group_name(file.file_name)
        elapsed_seconds = analyzed_at.timestamp() - file.modified_timestamp
        group_stats[group_name]["total"] += 1
        if elapsed_seconds <= ten_days:
            group_stats[group_name]["within_10"] += 1
        elif elapsed_seconds <= thirty_days:
            group_stats[group_name]["over_10_within_30"] += 1
        else:
            group_stats[group_name]["over_30"] += 1

    return group_stats


def _get_extension_group_name(file_name: str) -> str:
    extension = _get_extension_label(file_name)
    if extension.startswith("."):
        extension = extension[1:]
    for group_name, extensions in EXTENSION_GROUPS:
        if extension in extensions:
            return group_name
    return "기타"


def _get_top_level_paths_for_memo(memo: WorkMemo) -> set[str]:
    top_level_paths = {
        relative_path.split("/", 1)[0] for relative_path in memo.linked_folders
    }
    top_level_paths.update(
        relative_path.split("/", 1)[0] for relative_path in memo.linked_files
    )
    return top_level_paths


def _get_group_paths_for_memo(memo: WorkMemo) -> set[str]:
    # Each checked folder is its own group (not collapsed to top-level).
    groups: set[str] = set(memo.linked_folders)
    # Any linked file that isn't under an already-included folder gets grouped
    # by its top-level folder as a fallback.
    for relative_path in memo.linked_files:
        if not any(
            relative_path == f or relative_path.startswith(f"{f}/")
            for f in memo.linked_folders
        ):
            groups.add(relative_path.split("/", 1)[0])
    return groups


def _add_file_location_table(
    document: Document,
    files: list[AnalyzedFile],
) -> None:
    if not files:
        document.add_paragraph("없음")
        return

    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.style = "Table Grid"
    _set_row_cells(table.rows[0], ["파일위치", "수정일시"])
    for file in files:
        _set_row_cells(table.add_row(), [file.relative_path, file.modified_at])
    _set_file_location_table_column_widths(document, table)
    _align_file_location_table(table)


def _add_end_of_document_page(document: Document) -> None:
    document.add_section(WD_SECTION_START.NEW_PAGE)
    _set_section_vertical_alignment_center(document.sections[-1])
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("END OF DOCUMENT")
    run.font.size = Pt(25)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    brand_paragraph = document.add_paragraph()
    brand_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run = brand_paragraph.add_run(
        "본 문서는 [대표님의 인사담당자의 인수인계 10분]으로 작성되었습니다."
    )
    brand_run.font.size = Pt(9)
    brand_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _get_files_for_linked_folders(
    all_files: list[AnalyzedFile],
    linked_folders: list[str],
) -> list[AnalyzedFile]:
    return [
        file
        for file in all_files
        if any(
            file.relative_path == folder
            or file.relative_path.startswith(f"{folder}/")
            for folder in linked_folders
        )
    ]


def _get_recent_files_for_top_level_folder(
    all_files: list[AnalyzedFile],
    top_level_path: str,
    linked_files: list[str],
    files_by_relative_path: dict[str, AnalyzedFile] | None = None,
    linked_folders: list[str] | None = None,
) -> list[AnalyzedFile]:
    user_selected_files = _get_user_selected_files_for_top_level_folder(
        all_files,
        linked_files,
        top_level_path,
        files_by_relative_path,
    )
    user_selected_paths = {file.relative_path for file in user_selected_files}

    remaining_slots = max(PRIORITY_REVIEW_FILE_LIMIT - len(user_selected_files), 0)
    auto_candidates: list[AnalyzedFile] = []
    if remaining_slots:
        # Restrict auto-candidate search to the folders the user actually linked
        # under this top-level path, so files from unrelated sibling subfolders
        # are not included. Fall back to the top-level path itself only when no
        # linked folders exist under it (e.g. the user selected only files).
        scope_folders = (
            [
                f
                for f in (linked_folders or [])
                if f == top_level_path or f.startswith(f"{top_level_path}/")
            ]
            or [top_level_path]
        )
        # Cap the similarity-comparison pool to the most recently modified
        # candidates so _deduplicate_similar_file_names' O(n^2) comparison
        # doesn't scale with the full folder's file count.
        recent_candidates = sorted(
            _get_files_for_linked_folders(all_files, scope_folders),
            key=lambda file: file.modified_timestamp,
            reverse=True,
        )[:PRIORITY_REVIEW_DEDUPLICATION_CANDIDATE_LIMIT]
        unique_files = _deduplicate_similar_file_names(recent_candidates)
        sorted_candidates = sorted(
            unique_files,
            key=lambda file: file.modified_timestamp,
            reverse=True,
        )
        auto_candidates = [
            file
            for file in sorted_candidates
            if file.relative_path not in user_selected_paths
        ][:remaining_slots]

    return [*user_selected_files, *auto_candidates]


def _get_priority_review_files_for_top_level_folder(
    all_files: list[AnalyzedFile],
    top_level_path: str,
    linked_files: list[str],
    files_by_relative_path: dict[str, AnalyzedFile] | None = None,
    linked_folders: list[str] | None = None,
) -> list[AnalyzedFile]:
    user_selected_files = _get_user_selected_files_for_top_level_folder(
        all_files,
        linked_files,
        top_level_path,
        files_by_relative_path,
    )
    user_selected_paths = {file.relative_path for file in user_selected_files}

    remaining_slots = max(PRIORITY_REVIEW_FILE_LIMIT - len(user_selected_files), 0)
    auto_candidates: list[AnalyzedFile] = []
    if remaining_slots:
        # Restrict auto-candidate search to the folders the user actually linked
        # under this top-level path, so files from unrelated sibling subfolders
        # are not included. Fall back to the top-level path itself only when no
        # linked folders exist under it (e.g. the user selected only files).
        scope_folders = (
            [
                f
                for f in (linked_folders or [])
                if f == top_level_path or f.startswith(f"{top_level_path}/")
            ]
            or [top_level_path]
        )
        # Cap the similarity-comparison pool to the most recently modified
        # candidates so _deduplicate_similar_file_names' O(n^2) comparison
        # doesn't scale with the full folder's file count.
        recent_candidates = sorted(
            _get_files_for_linked_folders(all_files, scope_folders),
            key=lambda file: file.modified_timestamp,
            reverse=True,
        )[:PRIORITY_REVIEW_DEDUPLICATION_CANDIDATE_LIMIT]
        unique_files = _deduplicate_similar_file_names(recent_candidates)
        sorted_candidates = sorted(
            unique_files,
            key=lambda file: (
                _score_priority_review_file(file),
                file.modified_timestamp,
            ),
            reverse=True,
        )
        auto_candidates = [
            file
            for file in sorted_candidates
            if file.relative_path not in user_selected_paths
        ][:remaining_slots]

    return [*user_selected_files, *auto_candidates]


def _get_user_selected_files_for_top_level_folder(
    all_files: list[AnalyzedFile],
    linked_files: list[str],
    top_level_path: str,
    files_by_relative_path: dict[str, AnalyzedFile] | None = None,
) -> list[AnalyzedFile]:
    selected_relative_paths = [
        relative_path
        for relative_path in linked_files
        if relative_path == top_level_path
        or relative_path.startswith(f"{top_level_path}/")
    ]
    if files_by_relative_path is None:
        files_by_relative_path = {file.relative_path: file for file in all_files}
    return [
        files_by_relative_path[relative_path]
        for relative_path in selected_relative_paths
        if relative_path in files_by_relative_path
    ]


def _score_priority_review_file(file: AnalyzedFile) -> int:
    extension = _get_extension_label(file.file_name)
    score = 0

    if extension in DOCUMENT_EXTENSIONS:
        score += 40
    if file.size_bytes >= SMALL_FILE_SIZE_BYTES:
        score += 10
    if file.size_bytes == 0:
        score -= 50
    if file.file_name.startswith("~$"):
        score -= 40
    if extension in SUPPORTING_EXTENSIONS:
        score -= 20
    if file.file_name.startswith(".") or file.is_hidden_or_system:
        score -= 20

    return score


def _deduplicate_similar_file_names(
    files: list[AnalyzedFile],
) -> list[AnalyzedFile]:
    groups: list[list[AnalyzedFile]] = []
    for file in files:
        file_stem = Path(file.file_name).stem
        matched_group = None
        for group in groups:
            group_stem = Path(group[0].file_name).stem
            if SequenceMatcher(None, file_stem, group_stem).ratio() >= 0.7:
                matched_group = group
                break

        if matched_group is None:
            groups.append([file])
        else:
            matched_group.append(file)

    return [
        max(group, key=lambda file: file.modified_timestamp)
        for group in groups
    ]


def _set_default_font(document: Document) -> None:
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "맑은 고딕"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def _add_signature_table(document: Document) -> None:
    headers = ["구분", "부서", "직급/직책", "성명", "서명", "서명일"]

    section = document.sections[-1]
    available_width = section.page_width - section.left_margin - section.right_margin
    # available_width is in EMUs (1 inch = 914400 EMU = 1440 twips → 1 twip = 635 EMU)
    table_width_emu = int(available_width)  # 100% of body width
    table_width_dxa = table_width_emu // 635  # convert EMU → twips for XML

    # Column widths: 구분=12%, 부서=20%, 서명=15%; 직급/직책·성명·서명일 = 나머지 53% ÷ 3
    # Using ‰ (per-mille) arithmetic to avoid float rounding issues.
    _FIXED = [120, 200, 0, 0, 150, 0]  # ‰ for 구분, 부서, [flex], [flex], 서명, [flex]
    _remaining = 1000 - sum(_FIXED)    # 530 ‰ for the 3 flex columns
    _flex, _extra = divmod(_remaining, 3)
    _FIXED[2] = _flex + (1 if _extra > 0 else 0)
    _FIXED[3] = _flex + (1 if _extra > 1 else 0)
    _FIXED[5] = _flex
    col_widths_emu = [table_width_emu * t // 1000 for t in _FIXED]
    col_widths_emu[-1] = table_width_emu - sum(col_widths_emu[:-1])  # absorb rounding

    table = document.add_table(rows=3, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    tbl = table._tbl
    tblPr = tbl.tblPr
    for tag in ("w:jc", "w:tblW"):
        existing = tblPr.find(qn(tag))
        if existing is not None:
            tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(table_width_dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)

    for row in table.rows:
        for col_idx, cell in enumerate(row.cells):
            cell.width = col_widths_emu[col_idx]

    _set_row_cells(table.rows[0], headers)
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row, label in zip(table.rows[1:], ["인계자", "인수자"]):
        _set_row_cells(row, [label, "", "", "", "", ""])
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if col_idx == 0
                    else WD_ALIGN_PARAGRAPH.RIGHT
                )
        trPr = row._tr.get_or_add_trPr()
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), "600")
        trPr.append(trHeight)


def _add_separator(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph_properties = paragraph._p.get_or_add_pPr()
    paragraph_border = paragraph_properties.find(qn("w:pBdr"))
    if paragraph_border is None:
        paragraph_border = OxmlElement("w:pBdr")
        paragraph_properties.append(paragraph_border)

    bottom_border = OxmlElement("w:bottom")
    bottom_border.set(qn("w:val"), "single")
    bottom_border.set(qn("w:sz"), "6")
    bottom_border.set(qn("w:space"), "1")
    bottom_border.set(qn("w:color"), "auto")
    paragraph_border.append(bottom_border)


def _add_manual_heading(
    document: Document,
    text: str,
    font_size: int,
    space_before: int,
) -> None:
    paragraph = document.add_paragraph(style="Normal")
    _clear_numbering(paragraph)
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def _clear_numbering(paragraph: Any) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    numbering_properties = paragraph_properties.find(qn("w:numPr"))
    if numbering_properties is not None:
        paragraph_properties.remove(numbering_properties)


def _center_table_cells(table: Any) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _align_file_location_table(table: Any) -> None:
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in table.rows[1:]:
        location_cell = row.cells[0]
        modified_at_cell = row.cells[1]
        location_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        modified_at_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in location_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for paragraph in modified_at_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_file_location_table_column_widths(
    document: Document,
    table: Any,
) -> None:
    modified_at_width = Cm(4.7)
    section = document.sections[-1]
    available_width = section.page_width - section.left_margin - section.right_margin
    file_location_width = available_width - modified_at_width

    for row in table.rows:
        row.cells[0].width = file_location_width
        row.cells[1].width = modified_at_width


def _format_count(value: int) -> str:
    if value == 0:
        return "-"
    return f"{value}개"


def _set_section_vertical_alignment_center(section: Any) -> None:
    section_properties = section._sectPr
    vertical_alignment = section_properties.find(qn("w:vAlign"))
    if vertical_alignment is None:
        vertical_alignment = OxmlElement("w:vAlign")
        section_properties.append(vertical_alignment)
    vertical_alignment.set(qn("w:val"), "center")


def _get_extension_label(file_name: str) -> str:
    extension = Path(file_name).suffix
    if not extension:
        return "[no extension]"
    return extension.lower()


def _get_parent_relative_path(relative_path: str) -> str:
    if "/" not in relative_path:
        return "."
    return relative_path.rsplit("/", 1)[0]


def _get_folder_name(relative_path: str) -> str:
    return relative_path.rstrip("/").rsplit("/", 1)[-1]


def _add_key_value_table(document: Document, rows: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _set_row_cells(table.rows[0], ["항목", "값"])
    for key, value in rows:
        _set_row_cells(table.add_row(), [key, value])


def _set_row_cells(row: Any, values: list[Any]) -> None:
    for cell, value in zip(row.cells, values):
        cell.text = str(value)
